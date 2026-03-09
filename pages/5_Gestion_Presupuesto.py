# -*- coding: utf-8 -*-
# MÓDULO: GESTIÓN DE PRESUPUESTO — evidant Suite v3
# Mejoras: tipografía grande, homologación semántica, nombre programa en tablas,
#           pivot centrado legible, descomposición orden Resolución→Programa, gráficos con títulos

import sys, os, io, traceback, sqlite3, hashlib, unicodedata
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import ev_design

st.set_page_config(page_title="Gestión Presupuesto · Evidant", page_icon="💰", layout="wide")

ev_design.render(
    current="presupuesto",
    page_title="Gestión Presupuestaria",
    page_sub="Control de ejecución mensual · Marcos por CC · Dotación y análisis de costo",
    breadcrumb="Gestión Presupuestaria",
    icon="💰",
)
# ── ev_design ya inyectó el CSS Spotify maestro ──
# CSS complementario para tablas pivot y semáforos (no conflicta con ev_design)
st.markdown("""
<style>
/* ── Pivot table ──────────────────────────────────────── */
.ptbl{width:100%;border-collapse:collapse;font-size:.95rem;}
.ptbl th{background:#181818;color:var(--text2,#b3b3b3);padding:.65rem 1rem;
  border:1px solid rgba(255,255,255,.07);
  font-weight:700;text-align:center;font-size:.85rem;letter-spacing:.5px;text-transform:uppercase;}
.ptbl th.tleft{text-align:left;}
.ptbl td{padding:.55rem 1rem;border:1px solid rgba(255,255,255,.05);text-align:center;
  color:#d5e8f8;font-family:'JetBrains Mono',monospace;font-size:.93rem;}
.ptbl td.tleft{text-align:left;color:#b3b3b3;font-family:'Outfit',sans-serif;font-weight:600;}
.ptbl td.tprog{text-align:left;color:#535353;font-size:.78rem;padding-left:1.4rem;font-style:italic;}
.ptbl tr:nth-child(even){background:rgba(255,255,255,.02);}
.ptbl tr:hover{background:rgba(29,185,84,.05);}
.ptbl .tr-tot td{background:#1a1a1a;color:#1ed760!important;font-weight:800;border-top:2px solid #1db954;}
.ptbl .tr-tot td.tleft{font-family:'Outfit',sans-serif;}
.tc{color:#1ed760!important;font-weight:800;}
.lth{color:var(--text2,#b3b3b3)!important;font-size:.78rem;}
.ltd{color:var(--gold,#f59e0b)!important;font-size:.82rem;font-weight:600;}
.ltn{color:var(--green,#1db954)!important;font-size:.82rem;font-weight:700;}
/* ── Semáforo ── */
.ev-tag{display:inline-block;padding:.25rem .75rem;border-radius:20px;font-size:.8rem;font-weight:700;letter-spacing:.5px;}
.tag-ok{background:rgba(29,185,84,.14);color:#1ed760;border:1px solid rgba(29,185,84,.3);}
.tag-warn{background:rgba(245,158,11,.14);color:#f59e0b;border:1px solid rgba(245,158,11,.3);}
.tag-danger{background:rgba(239,68,68,.14);color:#ef4444;border:1px solid rgba(239,68,68,.3);}
</style>
""", unsafe_allow_html=True)

# ── DB ─────────────────────────────────────────────────────────────────────────
DB_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "presupuesto")
os.makedirs(DB_DIR, exist_ok=True)
DB_PATH = os.path.join(DB_DIR, "presupuesto.db")

def get_conn():
    c = sqlite3.connect(DB_PATH, check_same_thread=False, timeout=30)
    c.execute("PRAGMA journal_mode=WAL")
    c.execute("PRAGMA synchronous=NORMAL")
    return c

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR INFORME WORD — Reporte Ejecutivo Subtítulo 21
# ══════════════════════════════════════════════════════════════════════════════
def generar_informe_word(periodo_actual: str, periodo_anterior: str, logo_path: str = None, on_progress=None) -> io.BytesIO:
    """Informe ejecutivo Word — Reporte Mensual Subtítulo 21 D.A.P. SSMC.
    Incluye: portada con logo, resumen ejecutivo, análisis comparativo,
    análisis de programas, marcos presupuestarios, nuevas incorporaciones
    por modalidad (3 categorías CJ), bajas, alertas y recomendaciones."""

    MESES_ES = {1:"enero",2:"febrero",3:"marzo",4:"abril",5:"mayo",6:"junio",
                7:"julio",8:"agosto",9:"septiembre",10:"octubre",11:"noviembre",12:"diciembre"}

    def _lbl(p):
        try:
            pts = str(p).split("-")
            return f"{MESES_ES[int(pts[1])].capitalize()} {pts[0]}"
        except: return str(p)

    def _clp(v):
        try: return f"$ {int(round(float(v))):,}".replace(",",".")
        except: return "$ 0"

    def _pct(v, tot):
        try: return f"{float(v)/float(tot)*100:.1f}%" if float(tot) else "0%"
        except: return "—"

    def _n(v):
        try: return f"{int(v):,}".replace(",",".")
        except: return "0"

    # ── Homologación Calidad Jurídica → 3 categorías ────────────────────────
    def _cat_cj(v):
        s = str(v).strip().upper()
        if any(x in s for x in ["TITULAR","PLANTA"]): return "Titular de Planta"
        if any(x in s for x in ["CONTRAT"]): return "Contrata"
        if any(x in s for x in ["HONOR","SUMA ALZADA","A SUMA"]): return "Honorarios a Suma Alzada"
        return "Otro / Sin Clasificar"

    def _mov_lbl(v):
        s = str(v).strip().upper()
        if s in ("","NAN","NONE","NON"): return "Gasto Honorarios"
        if s == "NETO_ORIGEN": return "Gasto Neto Programa"
        if s == "DESCUENTO_DAP": return "Imputación D.A.P. Admin."
        return s

    _prog_step = [0]
    _PROG_TOTAL = 22
    def _prog(label=""):
        _prog_step[0] += 1
        if on_progress:
            on_progress(_prog_step[0], _PROG_TOTAL, label)

    # ── Cargar datos ────────────────────────────────────────────────────────
    conn = get_conn()
    df_pa      = pd.read_sql("SELECT * FROM personas_imputadas WHERE periodo=?",  conn, params=[periodo_actual])
    df_pn      = pd.read_sql("SELECT * FROM personas_imputadas WHERE periodo=?",  conn, params=[periodo_anterior])
    df_ia      = pd.read_sql("SELECT * FROM imputaciones WHERE periodo=?",         conn, params=[periodo_actual])
    df_in      = pd.read_sql("SELECT * FROM imputaciones WHERE periodo=?",         conn, params=[periodo_anterior])
    df_cc      = pd.read_sql("SELECT * FROM centros_costo",                        conn)
    df_all_imp = pd.read_sql("SELECT * FROM imputaciones ORDER BY periodo",        conn)
    conn.close()

    # Aplicar homologación CJ
    df_pa["cat_cj"] = df_pa["calidad_juridica"].apply(_cat_cj)
    df_pn["cat_cj"] = df_pn["calidad_juridica"].apply(_cat_cj)
    df_pa["_mov"]   = df_pa["tipo_movimiento"].apply(_mov_lbl)
    df_pn["_mov"]   = df_pn["tipo_movimiento"].apply(_mov_lbl)

    _prog("Datos cargados — construyendo portada...")
    lbl_act = _lbl(periodo_actual)
    lbl_ant = _lbl(periodo_anterior)

    # ── KPIs ────────────────────────────────────────────────────────────────
    hn_act  = df_pa["haber_neto"].sum()
    hn_ant  = df_pn["haber_neto"].sum()
    th_act  = df_pa["monto_total_haberes"].sum()
    th_ant  = df_pn["monto_total_haberes"].sum()
    ds_act  = df_pa["descuentos"].sum()
    ds_ant  = df_pn["descuentos"].sum()
    var_abs = hn_act - hn_ant
    var_pct = (var_abs / hn_ant * 100) if hn_ant else 0

    runs_act  = set(df_pa["run"].dropna().astype(str).str.strip().unique())
    runs_ant  = set(df_pn["run"].dropna().astype(str).str.strip().unique())
    runs_new  = runs_act - runs_ant
    runs_gone = runs_ant - runs_act

    df_new  = df_pa[df_pa["run"].astype(str).str.strip().isin(runs_new)].copy()
    df_gone = df_pn[df_pn["run"].astype(str).str.strip().isin(runs_gone)].copy()

    hon_act  = df_pa[df_pa["cat_cj"]=="Honorarios a Suma Alzada"]["haber_neto"].sum()
    con_act  = df_pa[df_pa["cat_cj"]=="Contrata"]["haber_neto"].sum()
    tit_act  = df_pa[df_pa["cat_cj"]=="Titular de Planta"]["haber_neto"].sum()
    hon_pct  = hon_act/hn_act*100 if hn_act else 0
    con_pct  = con_act/hn_act*100 if hn_act else 0
    tit_pct  = tit_act/hn_act*100 if hn_act else 0

    # ── Análisis de programas ───────────────────────────────────────────────
    prog_act = (df_ia.groupby("programa")
                .agg(haber_neto=("haber_neto","sum"),
                     total_haberes=("monto_total_haberes","sum"),
                     descuentos=("descuentos","sum"),
                     personas=("n_personas","sum"))
                .reset_index().sort_values("haber_neto", ascending=False))
    prog_ant = (df_in.groupby("programa")["haber_neto"]
                .sum().reset_index().rename(columns={"haber_neto":"hn_ant"}))
    prog_cmp = prog_act.merge(prog_ant, on="programa", how="left").fillna(0)
    prog_cmp["var_abs"] = prog_cmp["haber_neto"] - prog_cmp["hn_ant"]
    prog_cmp["var_pct"] = prog_cmp.apply(lambda r: r["var_abs"]/r["hn_ant"]*100 if r["hn_ant"] else 0, axis=1)
    prog_cmp["pct_tot"] = prog_cmp["haber_neto"] / hn_act * 100 if hn_act else 0
    top5_prog = prog_cmp.head(5)

    # ── Marcos presupuestarios ──────────────────────────────────────────────
    # La columna de clave en centros_costo se llama "numero" (no "cc")
    cc_map = {}
    if not df_cc.empty:
        for _, r in df_cc.iterrows():
            _key = str(r.get("numero", r.get("cc", ""))).strip()
            if _key:
                cc_map[_key] = {
                    "nombre": str(r.get("nombre","")).strip(),
                    "marco":  float(r.get("marco_clp", 0) or 0)
                }

    res_act = (df_ia.groupby("resolucion")
               .agg(haber_neto=("haber_neto","sum"), personas=("n_personas","sum"))
               .reset_index().sort_values("haber_neto", ascending=False))
    def _cc_lookup(x, field):
        """Busca en cc_map tolerando ceros a la izquierda y espacios."""
        s = str(x).strip()
        # Intento 1: clave exacta
        d = cc_map.get(s, None)
        if d is None:
            # Intento 2: sin ceros a la izquierda (numérico)
            try: d = cc_map.get(str(int(s)), None)
            except: pass
        if d is None:
            # Intento 3: con ceros a la izquierda (3 dígitos)
            try: d = cc_map.get(s.zfill(3), None)
            except: pass
        return d.get(field, "" if field=="nombre" else 0) if d else ("" if field=="nombre" else 0)

    res_act["nombre"]    = res_act["resolucion"].apply(lambda x: _cc_lookup(x, "nombre"))
    res_act["marco"]     = res_act["resolucion"].apply(lambda x: _cc_lookup(x, "marco"))
    res_act["ejec_pct"]  = res_act.apply(lambda r: r["haber_neto"]/r["marco"]*100 if r["marco"] else 0, axis=1)
    res_act["pct_gasto"] = res_act["haber_neto"] / hn_act * 100 if hn_act else 0
    top_cc = res_act.head(10)

    # ── Construir documento Word ────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)

    C_DARK  = RGBColor(0x1e, 0x3a, 0x5f)
    C_BLUE  = RGBColor(0x1a, 0x73, 0xe8)
    C_LIGHT = RGBColor(0x93, 0xc5, 0xfd)
    C_AMBR  = RGBColor(0xd9, 0x7b, 0x00)
    C_GRN   = RGBColor(0x1a, 0x7a, 0x3c)
    C_RED   = RGBColor(0xc0, 0x1f, 0x1f)
    C_GRAY  = RGBColor(0x47, 0x55, 0x69)
    C_WHITE = RGBColor(0xff, 0xff, 0xff)
    C_BLK   = RGBColor(0x0f, 0x11, 0x17)

    def _bg(cell, hex6):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex6); tcPr.append(shd)

    def _hdr_row(table, labels, bg="1e3a5f", sz=9):
        row = table.rows[0]
        for i, lbl in enumerate(labels):
            if i >= len(row.cells): break
            c = row.cells[i]; c.text = lbl
            _bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = C_WHITE; r.font.bold = True; r.font.size = Pt(sz)

    def _data_row(table, vals, sz=9.5, bold=False, bg_hex=None):
        row = table.add_row()
        for i, v in enumerate(vals):
            if i >= len(row.cells): break
            row.cells[i].text = str(v)
            if bg_hex: _bg(row.cells[i], bg_hex)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(sz); r.bold = bold
        return row

    def _total_row(table, ncols, label, amount, label_idx=0, amount_idx=-1):
        row = table.add_row()
        for i in range(min(ncols, len(row.cells))):
            _bg(row.cells[i], "1a3a5f")
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_WHITE
        row.cells[label_idx].text = label
        if amount_idx == -1: amount_idx = min(ncols-1, len(row.cells)-1)
        row.cells[amount_idx].text = amount
        for c in [row.cells[label_idx], row.cells[amount_idx]]:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.color.rgb = C_WHITE; r.font.size = Pt(9)

    def _h(lvl, txt, clr=None):
        p = doc.add_paragraph()
        run = p.add_run(txt); run.bold = True
        sizes = {1:16, 2:13, 3:11}; colors = {1:C_DARK, 2:C_BLUE, 3:C_BLUE}
        run.font.size = Pt(sizes.get(lvl, 11))
        run.font.color.rgb = clr or colors.get(lvl, C_BLUE)
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        return p

    def _body(txt, bold=False, italic=False, clr=None, sz=10.5):
        p = doc.add_paragraph()
        run = p.add_run(txt)
        run.font.size = Pt(sz); run.bold = bold; run.italic = italic
        if clr: run.font.color.rgb = clr
        p.paragraph_format.space_after = Pt(4)
        return p

    def _hr():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    # ────────────────────────────────────────────────────────────────────────
    # PORTADA
    # ────────────────────────────────────────────────────────────────────────
    # Buscar logo en assets/ (soporta PNG, JPG e ICO via Pillow)
    _ROOT   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    _ASSETS = os.path.join(_ROOT, "assets")
    _logo_candidates = [
        logo_path,
        os.path.join(_ASSETS, "logo.png"), os.path.join(_ASSETS, "logo_ssmc.png"),
        os.path.join(_ASSETS, "logo.jpg"), os.path.join(_ASSETS, "logo_ssmc.jpg"),
        os.path.join(_ASSETS, "logo.ico"),
        os.path.join(_ROOT, "logo.png"),
        os.path.join(_ROOT, "IMG_2426.ico"),
    ]
    _logo_file = next((p for p in _logo_candidates if p and os.path.exists(p)), None)

    # Convertir ICO → PNG en memoria para que python-docx lo acepte
    _logo_stream = None
    if _logo_file:
        try:
            from PIL import Image as _PILImage
            _img = _PILImage.open(_logo_file)
            # Seleccionar la variante de mayor resolución si es ICO multi-tamaño
            if hasattr(_img, 'ico') or _logo_file.lower().endswith(".ico"):
                sizes = getattr(_img, 'ico', None)
                if sizes:
                    _img.size  # trigger load
                try:
                    _img = _img.convert("RGBA")
                except Exception:
                    pass
            _buf_logo = io.BytesIO()
            _img.save(_buf_logo, format="PNG")
            _buf_logo.seek(0)
            _logo_stream = _buf_logo
        except Exception:
            # Si Pillow falla, intentar insertar directo (PNG/JPG nativo)
            _logo_stream = None

    # Encabezado de portada
    p_logo_row = doc.add_paragraph()
    p_logo_row.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _logo_stream or (_logo_file and not _logo_file.lower().endswith(".ico")):
        try:
            run_logo = p_logo_row.add_run()
            src = _logo_stream if _logo_stream else _logo_file
            run_logo.add_picture(src, width=Inches(1.6))
        except Exception:
            pass

    for sp in range(2): doc.add_paragraph()

    for txt, sz, clr, bold in [
        ("SERVICIO DE SALUD METROPOLITANO CENTRAL", 12, C_DARK, True),
        ("Dirección de Atención Primaria", 11, C_GRAY, True),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = clr; r.bold = bold

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REPORTE MENSUAL DE GASTOS"); r.font.size = Pt(24); r.bold = True; r.font.color.rgb = C_DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SUBTÍTULO 21 — CONTRATACIÓN DE PERSONAS"); r.font.size = Pt(16); r.bold = True; r.font.color.rgb = C_BLUE

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Análisis Comparativo:  {lbl_ant}  vs  {lbl_act}")
    r.font.size = Pt(13); r.bold = True; r.font.color.rgb = C_DARK

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Fecha de emisión: {datetime.now().day} de {MESES_ES[datetime.now().month]} de {datetime.now().year}")
    r.font.size = Pt(10); r.font.color.rgb = C_GRAY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Elaborado por: Unidad de Control de Gestión  ·  Evidant Suite")
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = C_GRAY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOCUMENTO DE USO INTERNO — CONFIDENCIAL")
    r.font.size = Pt(8); r.bold = True; r.font.color.rgb = C_RED

    doc.add_page_break()

    _prog("Portada generada — construyendo índice...")
    # ────────────────────────────────────────────────────────────────────────
    # ÍNDICE (manual)
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "ÍNDICE DE CONTENIDOS")
    indice = [
        ("1.", "Resumen Ejecutivo"),
        ("2.", "Indicadores Clave del Período"),
        ("3.", "Análisis Comparativo Mensual"),
        ("4.", "Análisis de Programas de Salud"),
        ("5.", "Marcos Presupuestarios — Ejecución por Resolución"),
        ("6.", "Nuevas Incorporaciones del Período"),
        ("7.", "Bajas del Período"),
        ("8.", "Análisis por Calidad Jurídica (3 categorías)"),
        ("9.", "Alertas y Riesgos Identificados"),
        ("10.","Recomendaciones de Control de Gestión"),
    ]
    for num, tit in indice:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num}  "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = C_BLUE
        r2 = p.add_run(tit); r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    _prog("Índice listo — generando secciones de análisis...")
    # ────────────────────────────────────────────────────────────────────────
    # 1. RESUMEN EJECUTIVO
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "1. RESUMEN EJECUTIVO")
    dir_var = "incremento" if var_abs >= 0 else "disminución"

    _body(
        f"El presente informe tiene por objeto entregar a la jefatura de la Dirección de Atención Primaria "
        f"(D.A.P.) del Servicio de Salud Metropolitano Central un análisis integral del comportamiento del "
        f"gasto en Subtítulo 21 — Contratación de Personas, con comparación entre los períodos de "
        f"{lbl_ant} y {lbl_act}, abarcando la totalidad de programas de salud administrados por la "
        f"Dirección.", sz=10.5
    )
    doc.add_paragraph()
    _body(
        f"Durante {lbl_act}, el gasto total en Haber Neto ascendió a {_clp(hn_act)}, registrando un "
        f"{dir_var} de {_clp(abs(var_abs))} ({abs(var_pct):.1f}%) respecto a {lbl_ant} ({_clp(hn_ant)}). "
        f"El Total de Haberes (gasto bruto antes de descuentos) fue de {_clp(th_act)}, con descuentos "
        f"por {_clp(ds_act)}, resultando en el Haber Neto señalado.", sz=10.5
    )
    doc.add_paragraph()
    _body(
        f"La dotación activa registrada en el período fue de {df_pa['run'].nunique():,} personas, "
        f"distribuidas entre {len(prog_act):,} programas de salud. Se registraron {len(runs_new)} nuevas "
        f"incorporaciones y {len(runs_gone)} personas sin imputación respecto al mes anterior. "
        f"La estructura del gasto por modalidad contractual muestra que los Honorarios a Suma Alzada "
        f"representan el {hon_pct:.1f}% del total ({_clp(hon_act)}), las Contratas el {con_pct:.1f}% "
        f"({_clp(con_act)}) y los Titulares de Planta el {tit_pct:.1f}% ({_clp(tit_act)}).", sz=10.5
    )
    doc.add_paragraph()
    if len(top5_prog) > 0:
        top1 = top5_prog.iloc[0]
        _body(
            f"El programa de mayor ejecución de gasto fue «{str(top1['programa'])[:70]}» con "
            f"{_clp(top1['haber_neto'])} ({top1['pct_tot']:.1f}% del total), seguido de "
            f"{'«' + str(top5_prog.iloc[1]['programa'])[:50] + '»' if len(top5_prog) > 1 else 'otros programas'} "
            f"con {_clp(top5_prog.iloc[1]['haber_neto']) if len(top5_prog) > 1 else ''}. "
            f"En términos acumulados, los 5 programas de mayor gasto concentran el "
            f"{top5_prog['pct_tot'].sum():.1f}% del gasto total del período.", sz=10.5
        )

    doc.add_page_break()

    _prog("Sección 1 — Indicadores clave...")
    # ────────────────────────────────────────────────────────────────────────
    # 2. INDICADORES CLAVE
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "2. INDICADORES CLAVE DEL PERÍODO")

    t_kpi = doc.add_table(rows=1, cols=4); t_kpi.style = "Table Grid"
    _hdr_row(t_kpi, ["Indicador Financiero", lbl_ant, lbl_act, "Variación"], sz=9)
    kpis = [
        ("Total Haberes (gasto bruto)",  _clp(th_ant), _clp(th_act),
         f"{'+' if th_act>=th_ant else ''}{((th_act-th_ant)/th_ant*100) if th_ant else 0:.1f}%"),
        ("Descuentos Previsionales",     _clp(ds_ant), _clp(ds_act), ""),
        ("Haber Neto Total",             _clp(hn_ant), _clp(hn_act),
         f"{'+' if var_abs>=0 else ''}{var_pct:.1f}%"),
        ("N° Personas en Nómina",        _n(df_pn["run"].nunique()), _n(df_pa["run"].nunique()),
         f"{df_pa['run'].nunique()-df_pn['run'].nunique():+d}"),
        ("N° Programas con Gasto",       _n(df_in["programa"].nunique()), _n(df_ia["programa"].nunique()), ""),
        ("Nuevas Incorporaciones",       "—", _n(len(runs_new)),  f"+{len(runs_new)}"),
        ("Bajas del Período",            "—", _n(len(runs_gone)), f"-{len(runs_gone)}"),
        ("Honorarios / Haber Neto",
         _pct(df_pn[df_pn["cat_cj"]=="Honorarios a Suma Alzada"]["haber_neto"].sum(), hn_ant),
         f"{hon_pct:.1f}%", ""),
        ("Contratas / Haber Neto",       "", f"{con_pct:.1f}%", ""),
        ("Titulares / Haber Neto",       "", f"{tit_pct:.1f}%", ""),
        ("Costo Promedio por Persona",
         _clp(hn_ant/df_pn["run"].nunique()) if df_pn["run"].nunique() else "$ 0",
         _clp(hn_act/df_pa["run"].nunique()) if df_pa["run"].nunique() else "$ 0", ""),
    ]
    for rd in kpis:
        r = _data_row(t_kpi, rd, sz=9.5)
        try:
            v3 = str(rd[3]).strip()
            if v3 and v3 not in ("", "—"):
                pv = float(v3.replace("+","").replace("%","").replace("$","").replace(".","").replace(",",".").strip())
                clr_v = C_RED if pv > 3 else (C_GRN if pv < -1 else C_AMBR)
                for p in r.cells[3].paragraphs:
                    for run in p.runs: run.font.color.rgb = clr_v
        except: pass

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 2 — Análisis comparativo...")
    # ────────────────────────────────────────────────────────────────────────
    # 3. ANÁLISIS COMPARATIVO MENSUAL
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "3. ANÁLISIS COMPARATIVO MENSUAL")

    _h(2, f"3.1  Gasto por Modalidad Contractual — {lbl_ant} vs {lbl_act}")
    _body(
        "La siguiente tabla desglosa el gasto por tipo de contrato, permitiendo identificar variaciones "
        "en la estructura de la dotación y su impacto financiero."
    )
    comp_cat = (df_pa.groupby("cat_cj")["haber_neto"]
                .agg(monto="sum", pers="nunique").reset_index()
                .merge(df_pn.groupby("cat_cj")["haber_neto"].agg(monto_ant="sum").reset_index(),
                       on="cat_cj", how="outer").fillna(0))
    comp_cat["var_$"] = comp_cat["monto"] - comp_cat["monto_ant"]
    comp_cat["var_%"] = comp_cat.apply(lambda r: r["var_$"]/r["monto_ant"]*100 if r["monto_ant"] else 0, axis=1)
    comp_cat["pct"]   = comp_cat["monto"] / hn_act * 100 if hn_act else 0
    comp_cat = comp_cat.sort_values("monto", ascending=False)

    t_cat = doc.add_table(rows=1, cols=7); t_cat.style = "Table Grid"
    _hdr_row(t_cat, ["Calidad Jurídica", f"HN {lbl_ant}", f"HN {lbl_act}", "Variación $",
                     "Var. %", "% del Total", "Personas"])
    for _, row in comp_cat.iterrows():
        r = _data_row(t_cat, [row["cat_cj"], _clp(row["monto_ant"]), _clp(row["monto"]),
                               _clp(row["var_$"]), f"{row['var_%']:.1f}%",
                               f"{row['pct']:.1f}%", int(row["pers"])], sz=9.5)
        for p in r.cells[4].paragraphs:
            for run in p.runs:
                run.font.color.rgb = C_RED if row["var_%"] > 3 else (C_GRN if row["var_%"] < -1 else C_GRAY)
    _total_row(t_cat, 7, "TOTAL", _clp(hn_act), 0, 2)

    doc.add_paragraph()
    _h(2, f"3.2  Gasto por Tipo de Movimiento Financiero")
    t_mv = doc.add_table(rows=1, cols=5); t_mv.style = "Table Grid"
    _hdr_row(t_mv, ["Tipo de Movimiento", f"Monto {lbl_ant}", f"Monto {lbl_act}", "Variación $", "Var. %"])
    comp_mv = (df_pa.groupby("_mov")["haber_neto"].sum().reset_index()
               .merge(df_pn.groupby("_mov")["haber_neto"].sum().reset_index().rename(columns={"haber_neto":"ant"}),
                      on="_mov", how="outer").fillna(0))
    comp_mv["var"] = comp_mv["haber_neto"] - comp_mv["ant"]
    comp_mv["vp"]  = comp_mv.apply(lambda r: r["var"]/r["ant"]*100 if r["ant"] else 0, axis=1)
    for _, row in comp_mv.iterrows():
        _data_row(t_mv, [row["_mov"], _clp(row["ant"]), _clp(row["haber_neto"]),
                         _clp(row["var"]), f"{row['vp']:.1f}%"], sz=9.5)
    _total_row(t_mv, 5, "TOTAL", _clp(hn_act), 0, 2)

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 3 — Análisis de programas...")
    # ────────────────────────────────────────────────────────────────────────
    # 4. ANÁLISIS DE PROGRAMAS DE SALUD
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "4. ANÁLISIS DE PROGRAMAS DE SALUD")
    _body(
        f"La D.A.P. gestiona {len(prog_act):,} programas de salud con imputación de gasto "
        f"en Subtítulo 21 durante {lbl_act}. El análisis de concentración del gasto revela que los "
        f"5 primeros programas representan el {top5_prog['pct_tot'].sum():.1f}% del total ejecutado, "
        f"lo que indica un patrón de gasto altamente concentrado que requiere atención preferente "
        f"en el control de ejecución presupuestaria."
    )
    doc.add_paragraph()

    if len(top5_prog) > 0 and top5_prog.iloc[0]["pct_tot"] > 30:
        _body(
            f"ANÁLISIS DE CONCENTRACIÓN: El programa «{str(top5_prog.iloc[0]['programa'])[:60]}» "
            f"concentra el {top5_prog.iloc[0]['pct_tot']:.1f}% del gasto total, lo que representa "
            f"una dependencia significativa. Desde la perspectiva del control de gestión, una "
            f"concentración superior al 30% en un único programa implica riesgo de reasignación "
            f"presupuestaria y menor flexibilidad ante requerimientos extraordinarios.",
            bold=False, clr=C_DARK
        )
        doc.add_paragraph()

    _h(2, "4.1  Todos los Programas — Ranking por Gasto Haber Neto")
    t_prog = doc.add_table(rows=1, cols=7); t_prog.style = "Table Grid"
    _hdr_row(t_prog, ["Programa", f"HN {lbl_ant}", f"HN {lbl_act}", "Total Haberes",
                       "Descuentos", "Variación $", "% del Total"])
    for rank, (_, row) in enumerate(prog_cmp.iterrows(), 1):
        bg = "f0f4ff" if rank <= 5 else None
        r = _data_row(t_prog, [
            str(row["programa"])[:55],
            _clp(row["hn_ant"]),
            _clp(row["haber_neto"]),
            _clp(row["total_haberes"]),
            _clp(row["descuentos"]),
            _clp(row["var_abs"]),
            f"{row['pct_tot']:.1f}%",
        ], sz=8.5, bg_hex=bg)
        for p in r.cells[5].paragraphs:
            for run in p.runs:
                run.font.color.rgb = C_RED if row["var_abs"] > 0 else C_GRN
    _total_row(t_prog, 7, "TOTAL", _clp(hn_act), 0, 2)

    doc.add_paragraph()
    _body(
        f"Nota: Los 5 programas destacados (fondo claro) concentran el {top5_prog['pct_tot'].sum():.1f}% "
        f"del gasto total. Se recomienda priorizar el seguimiento mensual de estos programas y "
        f"establecer alertas de ejecución en el módulo de control presupuestario.",
        italic=True, clr=C_GRAY, sz=9
    )

    doc.add_paragraph()
    _h(2, "4.2  Análisis Financiero de los Top 5 Programas")
    for i, (_, row) in enumerate(top5_prog.iterrows(), 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. {str(row['programa'])[:60]}  ")
        r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = C_BLUE
        r2 = p.add_run(
            f"Haber Neto: {_clp(row['haber_neto'])} ({row['pct_tot']:.1f}% del total). "
            f"Variación vs mes anterior: {'+' if row['var_abs']>=0 else ''}{_clp(row['var_abs'])} "
            f"({'+' if row['var_pct']>=0 else ''}{row['var_pct']:.1f}%). "
        )
        r2.font.size = Pt(10.5)
        trend = ("Presenta alza respecto al período anterior — revisar dotación autorizada." if row["var_abs"] > 0
                 else "Sin incremento relevante respecto al período anterior." if abs(row["var_abs"]) < abs(hn_act)*0.01
                 else "Registra baja, verificar imputaciones pendientes o bajas de dotación.")
        r3 = p.add_run(trend)
        r3.font.size = Pt(10); r3.italic = True
        r3.font.color.rgb = C_RED if row["var_abs"] > abs(hn_act)*0.02 else C_GRAY
        p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 4 — Marcos presupuestarios...")
    # ────────────────────────────────────────────────────────────────────────
    # 5. MARCOS PRESUPUESTARIOS
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "5. MARCOS PRESUPUESTARIOS — EJECUCIÓN POR RESOLUCIÓN / CC")
    _body(
        "La siguiente tabla presenta la ejecución presupuestaria por Centro de Costo (Resolución), "
        "contrastando el gasto imputado en Haber Neto con el marco presupuestario asignado. "
        "Los centros con ejecución superior al 80% del marco deben ser priorizados en el "
        "seguimiento mensual para evitar sobreejecución."
    )
    doc.add_paragraph()

    t_cc = doc.add_table(rows=1, cols=6); t_cc.style = "Table Grid"
    _hdr_row(t_cc, ["CC / Resolución", "Nombre", "Marco Asignado (CLP)",
                     "Haber Neto Ejecutado", "% Ejec.", "% del Gasto Total"])
    for _, row in top_cc.iterrows():
        ejec_pct_v = float(row["ejec_pct"])
        bg = "fff0f0" if ejec_pct_v >= 90 else ("fffbe6" if ejec_pct_v >= 75 else None)
        r = _data_row(t_cc, [
            str(row["resolucion"]),
            str(row["nombre"])[:40],
            _clp(row["marco"]) if row["marco"] > 0 else "Sin marco",
            _clp(row["haber_neto"]),
            f"{ejec_pct_v:.1f}%" if row["marco"] > 0 else "N/D",
            f"{row['pct_gasto']:.1f}%",
        ], sz=9, bg_hex=bg)
        for p in r.cells[4].paragraphs:
            for run in p.runs:
                run.font.color.rgb = C_RED if ejec_pct_v >= 90 else (C_AMBR if ejec_pct_v >= 75 else C_GRN)

    marcos_totales = df_cc["marco_clp"].sum() if not df_cc.empty and "marco_clp" in df_cc.columns else 0
    _total_row(t_cc, 6, "TOTAL GASTO", _clp(hn_act), 0, 3)

    doc.add_paragraph()
    _body("Semáforo: fondo rojo = ejecución ≥ 90%  |  fondo amarillo = ejecución 75–89%  |  sin color = < 75%",
          italic=True, clr=C_GRAY, sz=9)

    if marcos_totales > 0:
        ejec_global = hn_act/marcos_totales*100
        doc.add_paragraph()
        _body(
            f"EJECUCIÓN GLOBAL: El marco presupuestario total configurado en el sistema es "
            f"{_clp(marcos_totales)}, con una ejecución global del {ejec_global:.1f}% en {lbl_act}. "
            f"{'Se recomienda revisión urgente de la disponibilidad presupuestaria.' if ejec_global > 80 else 'La ejecución se mantiene dentro de rangos normales de gestión.'}",
            bold=True if ejec_global > 80 else False,
            clr=C_RED if ejec_global > 80 else C_GRN
        )

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 5 — Nuevas incorporaciones...")
    # ────────────────────────────────────────────────────────────────────────
    # 6. NUEVAS INCORPORACIONES
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "6. NUEVAS INCORPORACIONES DEL PERÍODO")
    _body(
        f"Se identificaron {len(runs_new)} personas que registran imputación de gasto en {lbl_act} "
        f"sin presencia en el período anterior ({lbl_ant}). El costo total de nuevas incorporaciones "
        f"asciende a {_clp(df_new['haber_neto'].sum())} en Haber Neto, representando el "
        f"{df_new['haber_neto'].sum()/hn_act*100:.1f}% del gasto total del período."
    )

    def _tabla_incorporaciones(df_sub, titulo, bg_hdr="1e3a5f"):
        _h(2, titulo)
        if df_sub.empty:
            _body("No se registran incorporaciones en esta modalidad.", italic=True, clr=C_GRAY); return
        df_d2 = df_sub.sort_values("haber_neto", ascending=False).drop_duplicates(subset=["run"])
        t = doc.add_table(rows=1, cols=8); t.style = "Table Grid"
        _hdr_row(t, ["RUT","Nombre","Calidad Jurídica","Programa","Unidad","Planilla","T. Haberes","Haber Neto"],
                 bg=bg_hdr, sz=8)
        for _, row in df_d2.iterrows():
            _data_row(t, [
                str(row.get("run","")).strip(),
                str(row.get("nombre","")).strip()[:28],
                str(row.get("calidad_juridica","")).strip()[:20],
                str(row.get("programa","")).strip()[:25],
                str(row.get("descripcion_unidad","")).strip()[:22],
                str(row.get("planilla_pago","")).strip(),
                _clp(row.get("monto_total_haberes",0)),
                _clp(row.get("haber_neto",0)),
            ], sz=8)
        _total_row(t, 8, f"TOTAL ({len(df_d2)} personas)", _clp(df_d2["haber_neto"].sum()), 0, 7)
        _body(f"Total {titulo}: {len(df_d2)} personas  ·  Haber Neto: {_clp(df_d2['haber_neto'].sum())}",
              bold=True, clr=C_BLUE)
        doc.add_paragraph()

    _tabla_incorporaciones(df_new[df_new["cat_cj"]=="Honorarios a Suma Alzada"],
                           "6.1  Honorarios a Suma Alzada — Nuevas Incorporaciones", bg_hdr="3a2800")
    _tabla_incorporaciones(df_new[df_new["cat_cj"]=="Contrata"],
                           "6.2  Contrata — Nuevas Incorporaciones", bg_hdr="1e3a5f")
    _tabla_incorporaciones(df_new[df_new["cat_cj"]=="Titular de Planta"],
                           "6.3  Titular de Planta — Nuevas Incorporaciones", bg_hdr="1a3a1a")

    doc.add_page_break()

    _prog("Sección 6 — Bajas del período...")
    # ────────────────────────────────────────────────────────────────────────
    # 7. BAJAS
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "7. BAJAS DEL PERÍODO")
    _body(
        f"Se identificaron {len(runs_gone)} personas que registraron imputación en {lbl_ant} y no "
        f"presentan imputación en {lbl_act}. El costo que representaron en el período anterior fue de "
        f"{_clp(df_gone['haber_neto'].sum())}. Estas situaciones pueden corresponder a término de "
        f"contrato, no renovación de honorarios, licencias médicas prolongadas o errores de imputación."
    )
    if not df_gone.empty:
        df_gd = df_gone.sort_values("haber_neto", ascending=False).drop_duplicates(subset=["run"])
        t_g = doc.add_table(rows=1, cols=7); t_g.style = "Table Grid"
        _hdr_row(t_g, ["RUT","Nombre","Calidad Jurídica","Tipo Contrato","Programa","Unidad",f"HN {lbl_ant}"],
                 bg="3a1a1a", sz=8.5)
        for _, row in df_gd.iterrows():
            _data_row(t_g, [
                str(row.get("run","")).strip(),
                str(row.get("nombre","")).strip()[:30],
                _cat_cj(row.get("calidad_juridica","")),
                str(row.get("tipo_contrato","")).strip(),
                str(row.get("programa","")).strip()[:25],
                str(row.get("descripcion_unidad","")).strip()[:22],
                _clp(row.get("haber_neto",0)),
            ], sz=8.5)
        _total_row(t_g, 7, f"TOTAL ({len(df_gd)} personas)", _clp(df_gd["haber_neto"].sum()), 0, 6)
    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 7 — Calidad jurídica...")
    # ────────────────────────────────────────────────────────────────────────
    # 8. ANÁLISIS POR CALIDAD JURÍDICA (3 categorías)
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "8. ANÁLISIS POR CALIDAD JURÍDICA — TRES CATEGORÍAS")
    _body(
        "La calidad jurídica define el vínculo contractual entre el trabajador/prestador y la institución. "
        "Para efectos del presente análisis se homologan todas las variantes registradas en el sistema "
        "a tres categorías canónicas: (1) Titular de Planta, (2) Contrata y (3) Honorarios a Suma Alzada. "
        "Esta homologación es fundamental para el análisis de riesgo laboral y previsional."
    )
    doc.add_paragraph()
    cj3 = (df_pa.groupby("cat_cj")
           .agg(personas=("run","nunique"),
                total_h=("monto_total_haberes","sum"),
                descuentos=("descuentos","sum"),
                haber_neto=("haber_neto","sum"))
           .reset_index().sort_values("haber_neto", ascending=False))
    cj3["prom"] = cj3.apply(lambda r: r["haber_neto"]/r["personas"] if r["personas"] else 0, axis=1)
    cj3["pct"]  = cj3["haber_neto"] / hn_act * 100 if hn_act else 0

    # Tabla resumen 3 categorías
    t_cj3 = doc.add_table(rows=1, cols=7); t_cj3.style = "Table Grid"
    _hdr_row(t_cj3, ["Calidad Jurídica","N° Personas","Total Haberes","Descuentos",
                      "Haber Neto","Promedio/Persona","% del Total"])
    for _, row in cj3.iterrows():
        _data_row(t_cj3, [row["cat_cj"], int(row["personas"]), _clp(row["total_h"]),
                           _clp(row["descuentos"]), _clp(row["haber_neto"]),
                           _clp(row["prom"]), f"{row['pct']:.1f}%"], sz=9.5)
    _total_row(t_cj3, 7, "TOTAL", _clp(hn_act), 0, 4)

    doc.add_paragraph()
    # Detalle por programa y categoría
    _h(2, "8.1  Gasto por Calidad Jurídica y Programa (Top 10 Programas)")
    cj_prog = (df_pa.groupby(["cat_cj","programa"])["haber_neto"]
               .sum().reset_index().sort_values("haber_neto", ascending=False).head(10))
    t_cjp = doc.add_table(rows=1, cols=4); t_cjp.style = "Table Grid"
    _hdr_row(t_cjp, ["Calidad Jurídica","Programa","Haber Neto","% del Total"])
    for _, row in cj_prog.iterrows():
        _data_row(t_cjp, [row["cat_cj"], str(row["programa"])[:55],
                           _clp(row["haber_neto"]), _pct(row["haber_neto"], hn_act)], sz=9)

    doc.add_paragraph()
    # Análisis narrativo
    _body(
        f"INTERPRETACIÓN FINANCIERA: La distribución entre modalidades contractuales tiene "
        f"implicancias directas sobre el riesgo institucional. Los {int(cj3[cj3['cat_cj']=='Honorarios a Suma Alzada']['personas'].sum()):,} "
        f"prestadores a honorarios ({hon_pct:.1f}% del gasto) no generan vínculo laboral formal, "
        f"pero una proporción elevada puede derivar en contingencias laborales ante fiscalizaciones "
        f"de la Dirección del Trabajo. Las {int(cj3[cj3['cat_cj']=='Contrata']['personas'].sum()):,} personas "
        f"en calidad de contrata ({con_pct:.1f}%) representan el segmento de mayor flexibilidad "
        f"presupuestaria, dado que sus contratos no son indefinidos por definición legal.",
        clr=C_DARK
    )

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 8 — Alertas y riesgos...")
    # ────────────────────────────────────────────────────────────────────────
    # 9. ALERTAS Y RIESGOS
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "9. ALERTAS Y RIESGOS IDENTIFICADOS")
    _body("Clasificación: [CRÍTICO] riesgo inmediato — acción urgente  |  "
          "[MODERADO] seguimiento requerido  |  [INFORMATIVO] para conocimiento de jefatura.",
          italic=True, clr=C_GRAY, sz=9)
    doc.add_paragraph()

    alerts = []
    if var_pct > 5:
        alerts.append(("CRÍTICO", C_RED,
            f"Incremento presupuestario de {var_pct:.1f}% ({_clp(abs(var_abs))}) respecto a {lbl_ant}. "
            f"Verificar disponibilidad de marco presupuestario restante para el ejercicio."))
    elif var_pct > 2:
        alerts.append(("MODERADO", C_AMBR,
            f"Incremento de {var_pct:.1f}% respecto al período anterior. Mantener seguimiento mensual."))
    elif var_pct < -5:
        alerts.append(("INFORMATIVO", C_GRN,
            f"Disminución del gasto en {abs(var_pct):.1f}%. Verificar imputaciones pendientes."))

    if hon_pct > 50:
        alerts.append(("CRÍTICO", C_RED,
            f"El {hon_pct:.1f}% del gasto corresponde a honorarios a suma alzada. "
            f"Riesgo legal elevado por posible subordinación laboral. Revisar con asesoría jurídica."))
    elif hon_pct > 35:
        alerts.append(("MODERADO", C_AMBR,
            f"El {hon_pct:.1f}% del gasto ({_clp(hon_act)}) corresponde a honorarios. "
            f"Evaluar regularización contractual de prestadores de larga data."))

    if len(runs_new) > 15:
        alerts.append(("MODERADO", C_AMBR,
            f"Volumen alto de nuevas incorporaciones: {len(runs_new)} personas. "
            f"Verificar resoluciones exentas y disponibilidad presupuestaria."))

    if len(runs_gone) > 8:
        alerts.append(("INFORMATIVO", C_BLUE,
            f"{len(runs_gone)} personas sin imputación en {lbl_act} vs {lbl_ant}. "
            f"Confirmar término formal de contrato o irregularidad en rendición."))

    if len(top_cc[top_cc["ejec_pct"] >= 90]) > 0:
        ccs_criticos = top_cc[top_cc["ejec_pct"] >= 90]["resolucion"].tolist()
        alerts.append(("CRÍTICO", C_RED,
            f"Los CC/Resoluciones {', '.join(str(x) for x in ccs_criticos[:5])} presentan ejecución "
            f"≥ 90% de su marco presupuestario. Riesgo de sobreejecución en el período."))

    if top5_prog.iloc[0]["pct_tot"] > 40 if len(top5_prog) > 0 else False:
        alerts.append(("MODERADO", C_AMBR,
            f"Alta concentración de gasto: el programa «{str(top5_prog.iloc[0]['programa'])[:50]}» "
            f"representa el {top5_prog.iloc[0]['pct_tot']:.1f}% del total."))

    if not alerts:
        alerts.append(("INFORMATIVO", C_GRN, "Sin alertas críticas identificadas en el período analizado."))

    for nivel, clr, texto in alerts:
        p = doc.add_paragraph()
        r1 = p.add_run(f"[{nivel}]  "); r1.bold = True; r1.font.color.rgb = clr; r1.font.size = Pt(11)
        r2 = p.add_run(texto); r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(7)

    doc.add_paragraph(); doc.add_page_break()

    _prog("Sección 9 — Recomendaciones...")
    # ────────────────────────────────────────────────────────────────────────
    # 10. RECOMENDACIONES
    # ────────────────────────────────────────────────────────────────────────
    _h(1, "10. RECOMENDACIONES DE CONTROL DE GESTIÓN")
    _body("Las siguientes recomendaciones se fundamentan en el análisis del período y tienen por "
          "objetivo fortalecer el control presupuestario y la gestión de la dotación de personas:")
    doc.add_paragraph()

    recs = [
        ("Control mensual de marcos presupuestarios",
         f"Implementar revisión mensual obligatoria de la ejecución de cada CC antes del cierre de "
         f"rendición, con semáforo de alerta al alcanzar el 75% del marco. Los {len(top_cc[top_cc['ejec_pct']>=75]) if 'ejec_pct' in top_cc.columns else 'N'} "
         f"CC que superan el 75% de ejecución deben ser priorizados."),
        ("Regularización de prestadores a honorarios",
         f"El {hon_pct:.1f}% del gasto en honorarios ({_clp(hon_act)}) expone a la institución a "
         f"contingencias laborales. Se recomienda levantar un catastro de prestadores con más de "
         f"12 meses continuos y evaluar regularización vía contrata en el proceso de formulación presupuestaria."),
        ("Validación presupuestaria previa a contrataciones",
         "Establecer visación electrónica de disponibilidad presupuestaria como requisito previo al "
         "ingreso de toda nueva contratación o renovación, utilizando el módulo Evidant como plataforma de validación."),
        ("Conciliación nómina-contratos",
         f"Las {len(runs_gone)} bajas identificadas deben ser verificadas contra el registro de contratos "
         f"vigentes en RRHH. Se recomienda proceso mensual de conciliación entre nómina de imputación "
         f"y contratos activos para detectar pagos a personas con contrato vencido o no renovado."),
        ("Análisis de concentración de programas",
         f"Los 5 programas con mayor gasto concentran el {top5_prog['pct_tot'].sum():.1f}% del total. "
         f"Se recomienda establecer revisión trimestral de la dotación de estos programas comparada "
         f"con la planificación original de sus respectivos convenios de transferencia."),
        ("Proyección anual y plan de contingencia",
         f"Basado en los {_clp(hn_act)} ejecutados en {lbl_act}, la proyección anual lineal alcanzaría "
         f"los {_clp(hn_act*12)}. Contrastar con presupuesto aprobado y definir medidas correctivas "
         f"si la proyección supera el 100% del presupuesto disponible."),
        ("Clasificación de movimientos financieros",
         "Regularizar los registros sin tipo de movimiento o con clasificación 'Gasto Honorarios' "
         "por ausencia de código. Coordinar con la Unidad de Remuneraciones la correcta codificación "
         "NETO_ORIGEN / DESCUENTO_DAP para asegurar la integridad de los reportes de gestión."),
    ]

    for i, (titulo, texto) in enumerate(recs, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. {titulo}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = C_BLUE
        r2 = p.add_run(texto); r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph(); doc.add_paragraph()
    p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_f.add_run(
        f"Reporte generado automáticamente por Evidant Suite  ·  Control de Gestión D.A.P. SSMC  ·  "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}  ·  Documento Confidencial"
    )
    r_f.font.size = Pt(8); r_f.italic = True; r_f.font.color.rgb = C_GRAY

    _prog("Generando anexos gráficos — Figura A.1...")
    # ────────────────────────────────────────────────────────────────────────
    # ANEXOS — FIGURAS ANALÍTICAS (estilo revista científica)
    # ────────────────────────────────────────────────────────────────────────
    import plotly.io as _pio

    doc.add_page_break()
    _h(1, "ANEXOS — SOPORTE GRÁFICO DEL ANÁLISIS")
    _body(
        "Los gráficos presentados a continuación constituyen el respaldo visual del informe. "
        "Cada figura ha sido producida a partir de los datos registrados en el sistema Evidant Suite "
        "para los períodos analizados. El formato de rotulación sigue los estándares de "
        "publicación analítica (título descriptivo, nota metodológica e implicancia de gestión), "
        "facilitando su uso en presentaciones y reportes de jefatura.",
        sz=10, italic=True, clr=C_GRAY
    )
    doc.add_paragraph()

    # ── Paleta y estilos para fondo blanco (imprimible) ───────────────────
    _PAL10W = ["#1a73e8","#34a853","#fbbc05","#ea4335","#9c27b0",
               "#00bcd4","#e65100","#607d8b","#e91e63","#795548"]
    _WB = dict(
        paper_bgcolor="white", plot_bgcolor="#f8fafc",
        font=dict(family="Arial, sans-serif", color="#334155", size=12),
        title_font=dict(size=15, color="#1e3a5f", family="Arial, sans-serif"),
        margin=dict(l=70, r=30, t=65, b=70),
    )
    def _wax(**kw):
        return dict(gridcolor="#e2e8f0", zerolinecolor="#cbd5e1",
                    tickfont=dict(size=11, color="#475569"),
                    title_font=dict(size=12, color="#334155"), **kw)

    def _fig2png(fig, w=1380, h=560):
        try:
            return _pio.to_image(fig, format="png", width=w, height=h, scale=2)
        except Exception:
            return None

    def _add_figure(num, title, img_bytes, desc, implication, fuente="Sistema Evidant Suite — Base datos Subtítulo 21 D.A.P. SSMC"):
        if img_bytes:
            try:
                _pi = doc.add_paragraph()
                _pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _pi.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(6.2))
            except Exception:
                _body("(No se pudo generar la imagen)", italic=True, clr=C_GRAY, sz=8)
        # Pie de figura — formato IEEE/revista
        _pc = doc.add_paragraph()
        _pc.paragraph_format.left_indent  = Cm(0.4)
        _pc.paragraph_format.right_indent = Cm(0.4)
        _pc.paragraph_format.space_before = Pt(4)
        _pc.paragraph_format.space_after  = Pt(3)
        _r1 = _pc.add_run(f"Figura {num}. "); _r1.bold = True; _r1.font.size = Pt(9.5); _r1.font.color.rgb = C_BLUE
        _r2 = _pc.add_run(f"{title}. "); _r2.bold = True; _r2.font.size = Pt(9.5)
        _r3 = _pc.add_run(desc); _r3.font.size = Pt(9); _r3.italic = True
        if implication:
            _pi2 = doc.add_paragraph()
            _pi2.paragraph_format.left_indent  = Cm(0.4)
            _pi2.paragraph_format.right_indent = Cm(0.4)
            _pi2.paragraph_format.space_after  = Pt(2)
            _ri1 = _pi2.add_run("Implicancia de gestión: ")
            _ri1.bold = True; _ri1.font.size = Pt(9); _ri1.font.color.rgb = C_DARK
            _ri2 = _pi2.add_run(implication)
            _ri2.font.size = Pt(9); _ri2.italic = True; _ri2.font.color.rgb = C_GRAY
        _pf2 = doc.add_paragraph()
        _pf2.paragraph_format.left_indent = Cm(0.4)
        _pf2.paragraph_format.space_after = Pt(12)
        _rf2 = _pf2.add_run(f"Fuente: {fuente}  ·  Período: {lbl_act}  ·  Elaboración: Unidad de Control de Gestión D.A.P.")
        _rf2.font.size = Pt(7.5); _rf2.italic = True; _rf2.font.color.rgb = C_GRAY
        # línea divisoria
        _phr = doc.add_paragraph()
        _phr.paragraph_format.space_before = Pt(2); _phr.paragraph_format.space_after = Pt(10)

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.1 — EJECUCIÓN PRESUPUESTARIA POR CENTRO DE COSTO
    # ══════════════════════════════════════════════════════════════════════
    _top_cc_fig = res_act.head(14).copy()
    _lbl_cc = _top_cc_fig.apply(
        lambda r: (str(r["nombre"])[:24] if str(r["nombre"]).strip() else str(r["resolucion"])), axis=1)
    _fa1 = go.Figure()
    if _top_cc_fig["marco"].sum() > 0:
        _fa1.add_bar(name="Marco Presupuestario", x=_lbl_cc, y=_top_cc_fig["marco"],
                     marker_color="rgba(26,115,232,0.18)",
                     marker_line_color="#1a73e8", marker_line_width=1.5)
    _ejec_clrs = ["#ef4444" if p >= 90 else "#f59e0b" if p >= 75 else "#34a853"
                  for p in _top_cc_fig["ejec_pct"]]
    _fa1.add_bar(name="Haber Neto Ejecutado", x=_lbl_cc, y=_top_cc_fig["haber_neto"],
                 marker_color=_ejec_clrs,
                 marker_line_color="rgba(0,0,0,0.15)", marker_line_width=1,
                 text=[f"${v/1e6:.1f}M" for v in _top_cc_fig["haber_neto"]],
                 textposition="outside", textfont=dict(size=10, color="#334155"))
    _fa1.update_layout(**{**_WB,
        "title": dict(text="Figura A.1 — Ejecución Presupuestaria por Centro de Costo vs Marco Asignado",
                      font=dict(size=14, color="#1e3a5f")),
        "margin": dict(l=60, r=30, t=65, b=120)},
        barmode="overlay", height=580,
        xaxis=_wax(title="Centro de Costo / Resolución", tickangle=-40),
        yaxis=_wax(title="Monto (CLP)"),
        legend=dict(orientation="h", y=-0.38, font=dict(size=11), bgcolor="white"))
    _fa1.update_xaxes(tickfont=dict(size=9))
    _img_a1 = _fig2png(_fa1, w=1380, h=580)
    _add_figure("A.1",
        "Ejecución Presupuestaria por Centro de Costo vs Marco Asignado",
        _img_a1,
        f"Comparación del Haber Neto ejecutado (barra coloreada) versus el marco presupuestario asignado "
        f"(barra azul traslúcida) para los principales Centros de Costo (CC) del período {lbl_act}. "
        f"Semáforo de color: verde < 75% de ejecución, amarillo 75–89%, rojo ≥ 90%.",
        f"El gráfico permite identificar de forma inmediata qué CC presentan riesgo de sobreejecución "
        f"(barra roja que se aproxima o supera el marco), requiriendo revisión urgente de disponibilidad "
        f"presupuestaria antes del cierre del período.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.2 — DISTRIBUCIÓN DEL GASTO POR PROGRAMA (DONUT)
    # ══════════════════════════════════════════════════════════════════════
    _ep2 = prog_act.head(10).copy()
    _ep2["pct_v"] = _ep2["haber_neto"] / _ep2["haber_neto"].sum() * 100
    _ep2["lbl_leg"] = _ep2["programa"].apply(lambda x: str(x)[:46])
    _pull2 = [0.07 if i == 0 else 0 for i in range(len(_ep2))]
    _txt2  = [f"{r['pct_v']:.1f}%" if r["pct_v"] >= 3 else "" for _, r in _ep2.iterrows()]
    _fa2 = go.Figure(go.Pie(
        labels=_ep2["lbl_leg"], values=_ep2["haber_neto"],
        hole=0.44, pull=_pull2,
        marker=dict(colors=_PAL10W[:len(_ep2)], line=dict(color="white", width=2.5)),
        text=_txt2, textinfo="text", textposition="inside",
        textfont=dict(size=12, color="white", family="Arial Black"),
        hovertemplate="<b>%{label}</b><br>HN: $%{value:,.0f}<br>%{percent}<extra></extra>",
        showlegend=True))
    _fa2.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.2 — Distribución del Haber Neto por Programa de Salud · Top 10 · {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=10, r=260, t=65, b=30)},
        height=560,
        legend=dict(font=dict(size=10, color="#334155"), orientation="v",
                    x=1.01, y=1.0, xanchor="left", yanchor="top",
                    bgcolor="white", bordercolor="#e2e8f0", borderwidth=1))
    _prog("Figura A.2 — Distribución por programa...")
    _img_a2 = _fig2png(_fa2, w=1380, h=560)
    _add_figure("A.2",
        f"Distribución del Haber Neto por Programa de Salud — Top 10 Programas ({lbl_act})",
        _img_a2,
        f"Gráfico de anillo (donut) que presenta la participación porcentual de cada programa de salud "
        f"en el gasto total de Haber Neto del período {lbl_act}. El segmento destacado corresponde al "
        f"programa de mayor concentración. Los porcentajes se muestran únicamente para segmentos ≥ 3%.",
        f"La concentración del gasto en pocos programas indica una estructura presupuestaria con baja "
        f"diversificación, lo que eleva el riesgo ante reducciones de convenio o reasignaciones. "
        f"Se recomienda revisar la planificación de dotación de los programas con mayor participación.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.3 — EVOLUCIÓN MENSUAL DEL GASTO (TODOS LOS PERÍODOS)
    # ══════════════════════════════════════════════════════════════════════
    if not df_all_imp.empty:
        _ev3 = (df_all_imp.groupby("periodo")
                .agg(hn=("haber_neto","sum"), thb=("monto_total_haberes","sum"),
                     ds=("descuentos","sum"))
                .reset_index().sort_values("periodo"))
        _ev3["lbl"] = _ev3["periodo"].apply(_lbl)
        _fa3 = go.Figure()
        _fa3.add_bar(name="Total Haberes (Gasto Bruto)", x=_ev3["lbl"], y=_ev3["thb"],
                     marker_color="rgba(26,115,232,0.45)",
                     marker_line_color="#1a73e8", marker_line_width=1.5)
        _fa3.add_bar(name="Descuentos Previsionales", x=_ev3["lbl"], y=_ev3["ds"],
                     marker_color="rgba(251,188,5,0.55)",
                     marker_line_color="#d97706", marker_line_width=1.2)
        _fa3.add_scatter(name="Haber Neto", x=_ev3["lbl"], y=_ev3["hn"],
                         mode="lines+markers+text",
                         text=[f"${v/1e6:.1f}M" for v in _ev3["hn"]],
                         textposition="top center", textfont=dict(size=11, color="#166534"),
                         line=dict(color="#15803d", width=3),
                         marker=dict(size=10, color="#15803d",
                                     line=dict(color="white", width=2)))
        _fa3.update_layout(**{**_WB,
            "title": dict(text="Figura A.3 — Evolución Mensual del Gasto Subtítulo 21: Total Haberes · Descuentos · Haber Neto",
                          font=dict(size=13, color="#1e3a5f")),
            "margin": dict(l=70, r=30, t=65, b=70)},
            barmode="group", height=520,
            xaxis=_wax(title="Período"), yaxis=_wax(title="Monto (CLP)"),
            legend=dict(orientation="h", y=-0.25, font=dict(size=12), bgcolor="white"))
        _prog("Figura A.3 — Evolución mensual...")
        _img_a3 = _fig2png(_fa3, w=1380, h=520)
        _add_figure("A.3",
            "Evolución Mensual del Gasto Subtítulo 21: Total Haberes, Descuentos y Haber Neto",
            _img_a3,
            f"Serie temporal que incluye todos los períodos importados al sistema. Las barras azules "
            f"representan el gasto bruto (Total Haberes) y las barras ámbar los descuentos previsionales. "
            f"La línea verde muestra el Haber Neto resultante (gasto líquido), con valores etiquetados en millones de CLP.",
            f"La tendencia del Haber Neto permite proyectar el gasto anual y detectar meses con comportamiento "
            f"atípico (alza o baja no planificada). Una brecha creciente entre Total Haberes y Haber Neto "
            f"señala mayor carga previsional en la dotación, relevante para el análisis actuarial.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.4 — DOTACIÓN Y COSTO POR MODALIDAD CONTRACTUAL
    # ══════════════════════════════════════════════════════════════════════
    _ta4 = (df_ia.groupby("tipo_contrato")
            .agg(personas=("n_personas","sum"), haber_neto=("haber_neto","sum"))
            .reset_index().sort_values("haber_neto", ascending=False))
    _ct4_clr = {"HONORARIOS": "#d97706", "REMUNERACIONES": "#1a73e8"}
    _fa4 = make_subplots(rows=1, cols=2,
                         subplot_titles=["N° Personas por Modalidad Contractual",
                                         "Haber Neto (CLP) por Modalidad Contractual"],
                         column_widths=[0.48, 0.52])
    for _, _rr4 in _ta4.iterrows():
        _c4 = _ct4_clr.get(str(_rr4["tipo_contrato"]).upper(), "#607d8b")
        _fa4.add_bar(row=1, col=1, x=[str(_rr4["tipo_contrato"])], y=[_rr4["personas"]],
                     marker_color=_c4, marker_line_color="rgba(0,0,0,0.15)", marker_line_width=1.5,
                     showlegend=False,
                     text=[f"{int(_rr4['personas']):,}"], textposition="outside",
                     textfont=dict(size=13, color="#334155"))
        _fa4.add_bar(row=1, col=2, x=[str(_rr4["tipo_contrato"])], y=[_rr4["haber_neto"]],
                     marker_color=_c4, marker_line_color="rgba(0,0,0,0.15)", marker_line_width=1.5,
                     showlegend=False,
                     text=[f"${_rr4['haber_neto']/1e6:.1f}M"], textposition="outside",
                     textfont=dict(size=12, color="#334155"))
    _fa4.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.4 — Dotación y Costo Laboral por Modalidad Contractual · {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=60, r=30, t=80, b=60)},
        height=500,
        xaxis=_wax(title="Modalidad"), yaxis=_wax(title="N° Personas"),
        xaxis2=_wax(title="Modalidad"), yaxis2=_wax(title="Haber Neto (CLP)"))
    _fa4.update_annotations(font_size=12, font_color="#1e3a5f")
    _prog("Figura A.4 — Dotación y costo...")
    _img_a4 = _fig2png(_fa4, w=1380, h=500)
    _add_figure("A.4",
        f"Dotación y Costo Laboral por Modalidad Contractual — Período {lbl_act}",
        _img_a4,
        f"Gráfico de barras duales que compara el volumen de dotación (N° de personas, panel izquierdo) "
        f"y el costo asociado en Haber Neto (panel derecho) según la modalidad contractual: "
        f"Honorarios (ámbar) vs Remuneraciones (azul). Datos del período actual.",
        f"La comparación permite verificar si existe proporcionalidad entre el número de personas "
        f"y el costo por modalidad. Una alta proporción de personas a honorarios con costo unitario "
        f"elevado puede indicar la necesidad de regularización contractual para optimizar la "
        f"estructura de gasto y reducir el riesgo legal de la institución.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.5 — GASTO POR UNIDAD DE DESEMPEÑO (TOP 15, APILADO HORIZONTAL)
    # ══════════════════════════════════════════════════════════════════════
    _ua5 = (df_ia.groupby(["descripcion_unidad","tipo_contrato"])
            .agg(costo=("haber_neto","sum")).reset_index()
            .sort_values("costo", ascending=False))
    _top_units = _ua5.groupby("descripcion_unidad")["costo"].sum().nlargest(15).index
    _ua5 = _ua5[_ua5["descripcion_unidad"].isin(_top_units)].copy()
    _ua5["descripcion_unidad"] = _ua5["descripcion_unidad"].str[:40]
    _fa5 = go.Figure()
    for _ct5, _clr5 in [("HONORARIOS","#d97706"), ("REMUNERACIONES","#1a73e8")]:
        _sub5 = _ua5[_ua5["tipo_contrato"]==_ct5]
        _all_u5 = _ua5["descripcion_unidad"].unique().tolist()
        _vals5 = [_sub5[_sub5["descripcion_unidad"]==u]["costo"].sum() for u in _all_u5]
        _fa5.add_bar(name=_ct5.capitalize(), y=_all_u5, x=_vals5, orientation="h",
                     marker_color=_clr5, marker_line_color="rgba(0,0,0,0.1)", marker_line_width=1)
    _fa5.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.5 — Haber Neto por Unidad de Desempeño · Top 15 · {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=230, r=30, t=65, b=70)},
        barmode="stack", height=620,
        xaxis=_wax(title="Haber Neto (CLP)"),
        yaxis=_wax(title=""),
        legend=dict(orientation="h", y=-0.12, font=dict(size=12), title_text="Modalidad: ", bgcolor="white"))
    _fa5.update_yaxes(autorange="reversed", tickfont=dict(size=10, color="#334155"))
    _prog("Figura A.5 — Unidades de desempeño...")
    _img_a5 = _fig2png(_fa5, w=1380, h=620)
    _add_figure("A.5",
        f"Haber Neto por Unidad de Desempeño — Top 15 Unidades · {lbl_act}",
        _img_a5,
        f"Barras horizontales apiladas que muestran el gasto en Haber Neto de las 15 unidades "
        f"de desempeño con mayor costo. Las barras están segmentadas por modalidad contractual "
        f"(Honorarios en ámbar, Remuneraciones en azul), permitiendo visualizar la composición "
        f"del gasto laboral al interior de cada unidad.",
        f"Las unidades con mayor proporción de honorarios (ámbar predominante) presentan mayor "
        f"exposición a riesgo laboral y menor estabilidad presupuestaria. Su identificación "
        f"permite priorizar auditorías de contratos y evaluar regularización contractual.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.6 — CALIDAD JURÍDICA: 3 CATEGORÍAS (GASTO + PROMEDIO POR PERSONA)
    # ══════════════════════════════════════════════════════════════════════
    _cj6 = cj3.copy()
    _fa6 = go.Figure()
    _fa6 = make_subplots(specs=[[{"secondary_y": True}]])
    _pal6 = ["#1a73e8","#d97706","#34a853","#ea4335"]
    for _i6, (_, _r6) in enumerate(_cj6.iterrows()):
        _fa6.add_bar(x=[str(_r6["cat_cj"])], y=[_r6["haber_neto"]],
                     name=str(_r6["cat_cj"]),
                     marker_color=_pal6[_i6 % len(_pal6)],
                     marker_line_color="rgba(0,0,0,0.15)", marker_line_width=1.5,
                     text=[f"${_r6['haber_neto']/1e6:.1f}M  ({_r6['pct']:.1f}%)"],
                     textposition="outside",
                     textfont=dict(size=11, color="#334155"),
                     secondary_y=False)
    _fa6.add_scatter(x=_cj6["cat_cj"].tolist(), y=_cj6["prom"].tolist(),
                     name="Costo Promedio/Persona",
                     mode="lines+markers+text",
                     text=[f"${v/1e3:.0f}K" for v in _cj6["prom"]],
                     textposition="top center", textfont=dict(size=10, color="#166534"),
                     line=dict(color="#15803d", width=2.5, dash="dot"),
                     marker=dict(size=10, color="#15803d", line=dict(color="white", width=2)),
                     secondary_y=True)
    _fa6.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.6 — Gasto Total y Costo Promedio por Persona según Calidad Jurídica · {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=80, r=80, t=65, b=80)},
        height=520,
        legend=dict(orientation="h", y=-0.25, font=dict(size=11), bgcolor="white"))
    _fa6.update_yaxes(title_text="Haber Neto Total (CLP)", gridcolor="#e2e8f0",
                      tickfont=dict(size=11), secondary_y=False)
    _fa6.update_yaxes(title_text="Costo Promedio por Persona (CLP)", gridcolor="rgba(0,0,0,0)",
                      tickfont=dict(size=11), secondary_y=True)
    _fa6.update_xaxes(tickfont=dict(size=12, color="#334155"))
    _prog("Figura A.6 — Calidad jurídica...")
    _img_a6 = _fig2png(_fa6, w=1380, h=520)
    _add_figure("A.6",
        f"Gasto Total en Haber Neto y Costo Promedio por Persona según Calidad Jurídica Homologada — {lbl_act}",
        _img_a6,
        f"Gráfico de barras con eje secundario de línea. Las barras muestran el gasto total en Haber Neto "
        f"para cada categoría jurídica homologada (Titular de Planta, Contrata, Honorarios a Suma Alzada). "
        f"La línea punteada verde (eje derecho) indica el costo promedio por persona en cada categoría.",
        f"El costo promedio por persona es un indicador de eficiencia del gasto laboral. Diferencias "
        f"significativas entre categorías pueden reflejar distintos niveles de calificación técnica o "
        f"jornadas parciales. Su seguimiento mensual permite detectar cambios estructurales en la dotación.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.7 — COMPARATIVO PERÍODO ANTERIOR VS ACTUAL POR MODALIDAD
    # ══════════════════════════════════════════════════════════════════════
    _cmp7 = comp_cat[comp_cat["cat_cj"] != "Otro / Sin Clasificar"].copy()
    _fa7  = go.Figure()
    _fa7.add_bar(name=lbl_ant, x=_cmp7["cat_cj"], y=_cmp7["monto_ant"],
                 marker_color="rgba(100,116,139,0.55)",
                 marker_line_color="#64748b", marker_line_width=1.5,
                 text=[f"${v/1e6:.1f}M" for v in _cmp7["monto_ant"]],
                 textposition="outside", textfont=dict(size=10, color="#475569"))
    _fa7.add_bar(name=lbl_act, x=_cmp7["cat_cj"], y=_cmp7["monto"],
                 marker_color="#1a73e8",
                 marker_line_color="#0d47a1", marker_line_width=1.5,
                 text=[f"${v/1e6:.1f}M" for v in _cmp7["monto"]],
                 textposition="outside", textfont=dict(size=10, color="#1e3a5f"))
    _fa7.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.7 — Comparativo del Gasto por Calidad Jurídica: {lbl_ant} vs {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=70, r=30, t=65, b=80)},
        barmode="group", height=500,
        xaxis=_wax(title="Calidad Jurídica"),
        yaxis=_wax(title="Haber Neto (CLP)"),
        legend=dict(orientation="h", y=-0.2, font=dict(size=13), bgcolor="white"))
    _prog("Figura A.7 — Comparativo de períodos...")
    _img_a7 = _fig2png(_fa7, w=1380, h=500)
    _add_figure("A.7",
        f"Análisis Comparativo del Gasto por Calidad Jurídica entre Períodos — {lbl_ant} vs {lbl_act}",
        _img_a7,
        f"Gráfico de barras agrupadas que contrasta el Haber Neto de cada categoría jurídica homologada "
        f"entre el período anterior ({lbl_ant}, barras grises) y el período actual ({lbl_act}, barras azules). "
        f"Los valores sobre las barras expresan el monto en millones de pesos.",
        f"Las variaciones entre períodos permiten identificar qué modalidad contractual está impulsando "
        f"el cambio en el gasto total. Un incremento en Honorarios no acompañado de aumento proporcional "
        f"en Contrata puede indicar contrataciones irregulares o expansión no programada de prestadores externos.")

    # ══════════════════════════════════════════════════════════════════════
    # FIGURA A.8 — CONCENTRACIÓN DE GASTO POR PROGRAMA (RANKING HORIZONTAL)
    # ══════════════════════════════════════════════════════════════════════
    _pc8 = prog_cmp.head(12).copy()
    _pc8["lbl8"] = _pc8["programa"].apply(lambda x: str(x)[:52])
    _pc8["pct_acum"] = _pc8["pct_tot"].cumsum()
    _fa8 = go.Figure()
    _fa8.add_bar(name=f"Haber Neto {lbl_act}", y=_pc8["lbl8"], x=_pc8["haber_neto"],
                 orientation="h",
                 marker_color=_PAL10W[:len(_pc8)],
                 marker_line_color="rgba(0,0,0,0.12)", marker_line_width=1,
                 text=[f"${v/1e6:.1f}M  ({p:.1f}%)" for v, p in zip(_pc8["haber_neto"], _pc8["pct_tot"])],
                 textposition="outside", textfont=dict(size=10, color="#334155"))
    _fa8.update_layout(**{**_WB,
        "title": dict(text=f"Figura A.8 — Ranking de Programas por Concentración de Gasto en Haber Neto · {lbl_act}",
                      font=dict(size=13, color="#1e3a5f")),
        "margin": dict(l=320, r=120, t=65, b=60)},
        height=620, showlegend=False,
        xaxis=_wax(title="Haber Neto (CLP)"), yaxis=_wax(title=""))
    _fa8.update_yaxes(autorange="reversed", tickfont=dict(size=9, color="#334155"))
    _prog("Figura A.8 — Ranking de programas...")
    _img_a8 = _fig2png(_fa8, w=1380, h=620)
    _add_figure("A.8",
        f"Ranking de Programas por Concentración de Gasto en Haber Neto — {lbl_act}",
        _img_a8,
        f"Barras horizontales ordenadas de mayor a menor gasto en Haber Neto, mostrando los "
        f"12 programas de salud de mayor impacto presupuestario en el período {lbl_act}. "
        f"Cada barra incluye el monto en millones de CLP y su participación porcentual sobre el total.",
        f"La concentración del gasto en los primeros programas del ranking es el principal determinante "
        f"del riesgo presupuestario de la D.A.P. Programas con alta participación deben contar con "
        f"revisión mensual de su dotación autorizada versus ejecutada, contrastando con los convenios "
        f"de transferencia correspondientes.")

    # Nota final de anexos
    doc.add_paragraph()
    _pnf = doc.add_paragraph()
    _pnf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rnf = _pnf.add_run(
        f"— Fin de Anexos —  ·  {len([_img_a1,_img_a2,_img_a3 if not df_all_imp.empty else None,_img_a4,_img_a5,_img_a6,_img_a7,_img_a8]):d} figuras analíticas  ·  "
        f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}  ·  Evidant Suite — D.A.P. SSMC"
    )
    _rnf.font.size = Pt(8); _rnf.italic = True; _rnf.font.color.rgb = C_GRAY

    _prog("Compilando documento final...")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def init_db():
    conn = get_conn()
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS centros_costo(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        numero TEXT UNIQUE NOT NULL, nombre TEXT NOT NULL,
        programa TEXT NOT NULL, marco_clp REAL DEFAULT 0,
        activo INTEGER DEFAULT 1, updated_at TEXT DEFAULT (datetime('now','localtime')));
    CREATE TABLE IF NOT EXISTS imputaciones(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        periodo TEXT NOT NULL, resolucion TEXT NOT NULL, programa TEXT NOT NULL,
        unidad TEXT, descripcion_unidad TEXT, tipo_contrato TEXT,
        tipo_movimiento TEXT, calidad_juridica TEXT,
        monto_total_haberes REAL DEFAULT 0, descuentos REAL DEFAULT 0,
        haber_neto REAL DEFAULT 0, n_personas INTEGER DEFAULT 0,
        lote_hash TEXT NOT NULL, created_at TEXT DEFAULT (datetime('now','localtime')));
    CREATE TABLE IF NOT EXISTS lotes_imputados(
        lote_hash TEXT PRIMARY KEY, periodo TEXT NOT NULL,
        n_registros INTEGER, created_at TEXT DEFAULT (datetime('now','localtime')));
    CREATE TABLE IF NOT EXISTS personas_imputadas(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        lote_hash TEXT NOT NULL, periodo TEXT NOT NULL,
        resolucion TEXT, programa TEXT, run TEXT, nombre TEXT,
        tipo_contrato TEXT, calidad_juridica TEXT, tipo_movimiento TEXT,
        unidad TEXT, descripcion_unidad TEXT,
        monto_total_haberes REAL DEFAULT 0, descuentos REAL DEFAULT 0,
        haber_neto REAL DEFAULT 0);
    """)
    for col, typ in [("monto_total_haberes","REAL DEFAULT 0"),("descuentos","REAL DEFAULT 0")]:
        try: conn.execute(f"ALTER TABLE imputaciones ADD COLUMN {col} {typ}")
        except: pass
    try: conn.execute("ALTER TABLE personas_imputadas ADD COLUMN planilla_pago TEXT")
    except: pass
    conn.commit(); conn.close()

init_db()

# ── Helpers ────────────────────────────────────────────────────────────────────
MESES = {1:"ENE",2:"FEB",3:"MAR",4:"ABR",5:"MAY",6:"JUN",
         7:"JUL",8:"AGO",9:"SEP",10:"OCT",11:"NOV",12:"DIC"}

# Homologación semántica: CONTRATOS=CONTRATADOS, TITULAR=TITULARES, etc.
_HOMO_CJ = {
    "CONTRATO":"CONTRATOS","CONTRATADO":"CONTRATOS",
    "CONTRATADOS":"CONTRATOS",
    "TITULAR":"TITULARES","TITULAR HORA":"TITULARES","TITULARES HORA":"TITULARES",
    "HONORARIO":"HONORARIOS","A HONORARIOS":"HONORARIOS",
}
def homo_cj(v):
    s = str(v).strip().upper() if pd.notna(v) and str(v).strip() not in ("","nan","None") else "S/D"
    return _HOMO_CJ.get(s, s)

def fmt_clp(v):
    try: return f"$ {int(round(float(v))):,}".replace(",",".")
    except: return "$ 0"

def fmt_n(v):
    try: return f"{int(round(float(v))):,}".replace(",",".")
    except: return "0"

def pct_bar(pct):
    pct = min(float(pct), 100)
    c = "#1db954" if pct<80 else "#f59e0b" if pct<95 else "#ef4444"
    return (f'<div style="background:#0f2845;border-radius:6px;height:11px;width:100%;margin-top:5px;">'+
            f'<div style="background:{c};width:{pct:.1f}%;height:11px;border-radius:6px;"></div></div>'+
            f'<div style="font-size:.88rem;color:{c};margin-top:3px;font-weight:700;">{pct:.1f}%</div>')

def semaforo(pct):
    if pct<80: return '<span class="ev-tag tag-ok">NORMAL</span>'
    if pct<95: return '<span class="ev-tag tag-warn">ALERTA</span>'
    return '<span class="ev-tag tag-danger">CRÍTICO</span>'

def parse_num(v):
    try:
        if isinstance(v,(int,float)): return float(v)
        return float(str(v).replace(".","").replace(",",".").strip())
    except: return 0.0

def per_label(p):
    try: y,m=p.split("-"); return f"{MESES.get(int(m),m)} {y}"
    except: return str(p)

def _norm(s):
    """Normaliza string: minúsculas, sin tildes, sin espacios extra."""
    s = str(s).lower().strip()
    return ''.join(c for c in unicodedata.normalize('NFD', s)
                   if unicodedata.category(c) != 'Mn')

def _find_col(df, *candidates):
    """
    Encuentra la columna del DataFrame que mejor coincide con alguno de los
    candidatos, ignorando tildes, mayúsculas y variaciones de longitud.
    Primero busca coincidencia exacta normalizada, luego parcial.
    """
    norm_map = {_norm(c): c for c in df.columns}
    for cand in candidates:
        nc = _norm(cand)
        if nc in norm_map:
            return norm_map[nc]
    # búsqueda parcial: el candidato está contenido en el nombre de la columna
    for cand in candidates:
        nc = _norm(cand)
        for nk, orig in norm_map.items():
            if nc in nk:
                return orig
    return None

# Plotly base — design system evidant
PB = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(17,21,32,0.8)",
    font=dict(family="JetBrains Mono, monospace", color="#94a3b8", size=13),
    title_font=dict(size=16, color="#f1f5f9", family="Cabinet Grotesk, sans-serif"),
    margin=dict(l=20, r=20, t=55, b=20),
)
def ax(**kw): return dict(
    gridcolor="rgba(255,255,255,.06)",
    zerolinecolor="rgba(255,255,255,.06)",
    tickfont=dict(size=12, color="#475569"),
    title_font=dict(size=13, color="#94a3b8"),
    **kw)

# ── DB helpers ─────────────────────────────────────────────────────────────────
def load_centros():
    conn=get_conn(); df=pd.read_sql("SELECT * FROM centros_costo WHERE activo=1 ORDER BY CAST(numero AS INTEGER),numero",conn); conn.close(); return df

def save_centro(num,nom,prog,marco):
    conn=get_conn()
    conn.execute("""INSERT INTO centros_costo(numero,nombre,programa,marco_clp) VALUES(?,?,?,?)
        ON CONFLICT(numero) DO UPDATE SET nombre=excluded.nombre,programa=excluded.programa,
        marco_clp=excluded.marco_clp,updated_at=datetime('now','localtime')""",
        (str(num).strip(),str(nom).strip(),str(prog).strip(),float(marco)))
    conn.commit(); conn.close()

def del_centro(num):
    conn=get_conn(); conn.execute("UPDATE centros_costo SET activo=0 WHERE numero=?",(num,)); conn.commit(); conn.close()

def load_imp():
    conn=get_conn(); df=pd.read_sql("SELECT * FROM imputaciones",conn); conn.close()
    if not df.empty and "calidad_juridica" in df.columns:
        df["calidad_juridica"] = df["calidad_juridica"].apply(homo_cj)
    return df

def load_lotes():
    conn=get_conn(); df=pd.read_sql("SELECT * FROM lotes_imputados ORDER BY periodo DESC",conn); conn.close(); return df

def lote_hash(df,periodo):
    k=periodo+str(len(df))+str(df.get("Total_Haberes_Netos",pd.Series()).sum())
    return hashlib.md5(k.encode()).hexdigest()[:16]

def build_prog_map(df_imp, df_cc):
    """Devuelve {resolucion: nombre_programa_formal_canonico}"""
    m = {}
    if not df_imp.empty:
        modal = (df_imp.groupby("resolucion")["programa"]
                 .agg(lambda s: s.mode()[0] if not s.mode().empty else s.iloc[0])
                 .to_dict())
        m.update(modal)
    if not df_cc.empty:
        for _, r in df_cc.iterrows():
            if str(r["numero"]).strip():
                m[str(r["numero"]).strip()] = str(r["programa"]).strip()
    return m

# ── Sidebar ────────────────────────────────────────────────────────────────────


tab_dash, tab_pivot, tab_import, tab_cc, tab_det = st.tabs([
    "📊  Dashboard", "📅  Estado Mensual", "📥  Imputar Rendiciones",
    "⚙️  Centros de Costo", "🔍  Análisis Detallado"])

# ══════════════════════════════════════════════════════════════════════════════
# CENTROS DE COSTO
# ══════════════════════════════════════════════════════════════════════════════
with tab_cc:
    st.markdown("### ⚙️ Configuración de Centros de Costo")
    df_cc  = load_centros(); df_ref = load_imp()
    prog_list = sorted(df_ref["programa"].dropna().unique().tolist()) if not df_ref.empty else []
    res_list  = sorted(df_ref["resolucion"].dropna().unique().tolist()) if not df_ref.empty else []
    cf, ct = st.columns([1, 1.8], gap="large")
    with cf:
        st.markdown("**Agregar / Actualizar Centro de Costo**")
        with st.form("form_cc", clear_on_submit=True):
            if res_list:
                r_opts = ["✏️ Ingresar manualmente..."] + res_list
                r_sel  = st.selectbox("N° Resolución (CC)", r_opts)
                n_num  = st.text_input("Resolución manual", placeholder="ej: 53") if r_sel=="✏️ Ingresar manualmente..." else r_sel
            else:
                n_num = st.text_input("N° Resolución (CC)", placeholder="ej: 53")
                st.caption("Importa rendiciones para cargar resoluciones.")
            n_nombre = st.text_input("Nombre formal", placeholder="ej: CESFAM CHUCHUNCO")
            if prog_list:
                p_opts = ["✏️ Ingresar manualmente..."] + prog_list
                p_sel  = st.selectbox("Programa asociado", p_opts)
                n_prog = st.text_input("Programa manual", placeholder="ej: APOYO A LA GESTION...") if p_sel=="✏️ Ingresar manualmente..." else p_sel
            else:
                n_prog = st.text_input("Programa asociado", placeholder="ej: APOYO A LA GESTION...")
            n_marco = st.number_input("Marco presupuestario anual (CLP)", min_value=0, step=1_000_000, value=0, format="%d")
            ok = st.form_submit_button("💾 Guardar Centro de Costo", use_container_width=True)
        if ok:
            if not str(n_num).strip() or not str(n_nombre).strip() or not str(n_prog).strip():
                st.error("Completa todos los campos.")
            else:
                save_centro(n_num, n_nombre, n_prog, n_marco); st.success(f"✅ CC **{n_num}** guardado."); st.rerun()
        if not df_cc.empty:
            st.markdown("---")
            opts = {f"{r.numero} — {r.nombre}": r.numero for r in df_cc.itertuples()}
            sel_d = st.selectbox("Eliminar CC", list(opts.keys()))
            if st.button("🗑️ Eliminar", use_container_width=True):
                del_centro(opts[sel_d]); st.success("Eliminado."); st.rerun()

        # ── Carga Masiva ────────────────────────────────────────────────────
        st.markdown("---")
        st.markdown("**📤 Carga Masiva desde Excel**")

        # Generar plantilla descargable
        _tpl_cols = ["N° Resolución", "Nombre Formal", "Programa", "Marco Presupuestario Anual (CLP)"]
        _tpl_data = [
            ["53", "CESFAM EJEMPLO NORTE", "APOYO A LA GESTION EN LOS ESTABLECIMIENTOS", 250_000_000],
            ["54", "CESFAM EJEMPLO SUR",   "REFORZAMIENTO APS",                           180_000_000],
        ]
        _tpl_df  = pd.DataFrame(_tpl_data, columns=_tpl_cols)
        _tpl_buf = io.BytesIO()
        with pd.ExcelWriter(_tpl_buf, engine="openpyxl") as _w:
            _tpl_df.to_excel(_w, index=False, sheet_name="Centros de Costo")
            ws = _w.sheets["Centros de Costo"]
            ws.column_dimensions["A"].width = 16
            ws.column_dimensions["B"].width = 32
            ws.column_dimensions["C"].width = 52
            ws.column_dimensions["D"].width = 30
        _tpl_buf.seek(0)
        st.download_button(
            "⬇️ Descargar formato de carga masiva",
            data=_tpl_buf.read(),
            file_name="formato_carga_centros_costo.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

        up_cc = st.file_uploader("Cargar Excel de centros de costo", type=["xlsx"], key="up_cc_masivo")
        if up_cc:
            try:
                _df_up = pd.read_excel(io.BytesIO(up_cc.read()), engine="openpyxl")
                _df_up.columns = [str(c).strip() for c in _df_up.columns]

                # Mapear columnas flexiblemente
                def _fc(df, *names):
                    for n in names:
                        if n in df.columns: return n
                        cands = [c for c in df.columns if n.lower() in c.lower()]
                        if cands: return cands[0]
                    return None

                _c_num  = _fc(_df_up, "N° Resolución", "Numero", "Resolución", "CC", "N°")
                _c_nom  = _fc(_df_up, "Nombre Formal", "Nombre")
                _c_prog = _fc(_df_up, "Programa")
                _c_mar  = _fc(_df_up, "Marco Presupuestario Anual (CLP)", "Marco", "marco_clp")

                _faltantes = [n for n, c in [("N° Resolución", _c_num), ("Nombre", _c_nom),
                                              ("Programa", _c_prog), ("Marco", _c_mar)] if not c]
                if _faltantes:
                    st.error(f"❌ Columnas no encontradas: {_faltantes}. Usa el formato descargable.")
                else:
                    _df_prev = _df_up[[_c_num, _c_nom, _c_prog, _c_mar]].dropna(subset=[_c_num]).copy()
                    _df_prev.columns = ["N° Res.", "Nombre", "Programa", "Marco CLP"]
                    _df_prev["Marco CLP"] = _df_prev["Marco CLP"].apply(lambda v: fmt_clp(parse_num(v)))
                    st.caption(f"Vista previa — {len(_df_prev)} registros detectados:")
                    st.markdown(ev_design.ev_table_html(_df_prev), unsafe_allow_html=True)

                    if st.button("⚡ Importar todos los centros de costo", type="primary", use_container_width=True):
                        _ok = _err = 0
                        for _, _r in _df_up.iterrows():
                            _num = str(_r[_c_num]).strip()
                            _nom = str(_r[_c_nom]).strip() if _c_nom else ""
                            _prg = str(_r[_c_prog]).strip() if _c_prog else ""
                            _mar = parse_num(_r[_c_mar]) if _c_mar else 0.0
                            if not _num or _num in ("nan", "None", ""):
                                continue
                            try:
                                save_centro(_num, _nom or _num, _prg or "SIN PROGRAMA", _mar)
                                _ok += 1
                            except Exception:
                                _err += 1
                        st.success(f"✅ {_ok} centros importados correctamente." + (f" ⚠️ {_err} errores." if _err else ""))
                        st.rerun()
            except Exception as _e:
                st.error(f"Error leyendo el archivo: {_e}")

    with ct:
        if df_cc.empty: st.info("No hay centros de costo. Agrega el primero.")
        else:
            ds = df_cc[["numero","nombre","programa","marco_clp"]].copy()
            ds.columns = ["N° Res.","Nombre","Programa","Marco CLP"]
            ds["Marco CLP"] = ds["Marco CLP"].apply(fmt_clp)
            st.markdown(ev_design.ev_table_html(ds), unsafe_allow_html=True)
        if res_list and not df_cc.empty:
            conf = set(df_cc["numero"].astype(str).tolist())
            sin  = [r for r in res_list if r not in conf]
            if sin: st.warning(f"⚠️ {len(sin)} resoluciones sin marco: {', '.join(sin[:10])}")

# ══════════════════════════════════════════════════════════════════════════════
# IMPUTAR RENDICIONES
# ══════════════════════════════════════════════════════════════════════════════
with tab_import:
    st.markdown("### 📥 Imputar Gastos desde Rendiciones")
    cu, ch = st.columns([1.3, 1], gap="large")
    with cu:
        up = st.file_uploader("4. RENDICIONES.xlsx", type=["xlsx"], key="rend_up")
        if up:
            try:
                rb = up.read()
                df_r = pd.read_excel(io.BytesIO(rb), sheet_name="RENDICIONES", engine="openpyxl")
                st.success(f"✅ **{len(df_r):,} registros** leídos.")
                ms = pd.to_numeric(df_r.get("Mes Rendicion", pd.Series(dtype=float)), errors="coerce").dropna()
                ys = pd.to_numeric(df_r.get("Anio Devengo",  pd.Series(dtype=float)), errors="coerce").dropna()
                md = int(ms.mode()[0]) if not ms.empty else datetime.now().month
                yd = int(ys.mode()[0]) if not ys.empty else datetime.now().year
                cm2,ca2 = st.columns(2)
                mes_s = cm2.selectbox("Mes", list(MESES.keys()), index=md-1, format_func=lambda x: f"{x:02d} — {MESES[x]}")
                ano_s = ca2.number_input("Año", min_value=2020, max_value=2035, value=yd)
                periodo = f"{ano_s}-{mes_s:02d}"
                h = lote_hash(df_r, periodo); lts = load_lotes()
                ya = not lts.empty and h in lts["lote_hash"].values
                if ya: st.warning(f"⚠️ **{per_label(periodo)}** ya fue imputado.")
                # ── detectar columnas con _find_col (tolerante a tildes) ──
                _cc_res = _find_col(df_r, "Resolucion","Resolución","Resolucion")
                _cc_prg = _find_col(df_r, "Programa")
                _cc_uni = _find_col(df_r, "Unidad")
                _cc_dun = _find_col(df_r, "Descripcion Unidad","Descripción Unidad")
                _cc_tc  = _find_col(df_r, "Tipo de Contrato")
                _cc_tm  = _find_col(df_r, "Tipo de Movimiento")
                _cc_cj  = _find_col(df_r, "Calidad Juridica","Calidad Jurídica")
                _cc_hn  = _find_col(df_r, "Total_Haberes_Netos","Total Haberes Netos")
                _cc_mth = _find_col(df_r, "Monto (Total Haberes)","Monto Total Haberes")
                _cc_dsc = _find_col(df_r, "Descuentos")
                _cc_run = _find_col(df_r, "RUN","RUT")
                _cc_nom = _find_col(df_r, "Proveedor o Prestador","Nombre","Prestador")
                _cc_plp = _find_col(df_r, "Planilla de Pago","Planilla Pago","Planilla")

                # Vista previa con columnas reales detectadas
                _pc_real = [c for c in [_cc_res,_cc_prg,_cc_tc,_cc_mth,_cc_dsc,_cc_hn] if c]
                st.caption(f"Columnas detectadas → Resolución:`{_cc_res}` | Programa:`{_cc_prg}` | Nombre:`{_cc_nom}` | Descuentos:`{_cc_dsc}`")
                if _pc_real:
                    st.markdown(ev_design.ev_table_html(df_r[_pc_real].head(5)), unsafe_allow_html=True)

                if st.button("⚡ IMPUTAR GASTOS", use_container_width=True, disabled=ya, type="primary"):
                    with st.status("⚡ Imputando gastos...", expanded=True) as _imp_status:
                      _pbar_imp = st.progress(0, text="Iniciando imputación...")
                    try:
                        df_r["_n"] = (df_r[_cc_hn].apply(parse_num)  if _cc_hn  else pd.Series(0, index=df_r.index))
                        df_r["_t"] = (df_r[_cc_mth].apply(parse_num) if _cc_mth else pd.Series(0, index=df_r.index))
                        df_r["_d"] = (df_r[_cc_dsc].apply(parse_num) if _cc_dsc else pd.Series(0, index=df_r.index))
                        df_r["_u"] = (df_r[_cc_run].astype(str)       if _cc_run else pd.Series("", index=df_r.index))

                        # Renombrar columnas clave a nombres canónicos para simplificar groupby
                        _rename = {}
                        if _cc_res and _cc_res != "Resolucion": _rename[_cc_res] = "Resolucion"
                        if _cc_prg and _cc_prg != "Programa":   _rename[_cc_prg] = "Programa"
                        if _cc_uni and _cc_uni != "Unidad":     _rename[_cc_uni] = "Unidad"
                        if _cc_dun and _cc_dun != "Descripcion Unidad": _rename[_cc_dun] = "Descripcion Unidad"
                        if _cc_tc  and _cc_tc  != "Tipo de Contrato":   _rename[_cc_tc]  = "Tipo de Contrato"
                        if _cc_tm  and _cc_tm  != "Tipo de Movimiento": _rename[_cc_tm]  = "Tipo de Movimiento"
                        if _cc_cj  and _cc_cj  != "Calidad Juridica":   _rename[_cc_cj]  = "Calidad Juridica"
                        df_r = df_r.rename(columns=_rename)

                        grp = [c for c in ["Resolucion","Programa","Descripcion Unidad","Unidad",
                                           "Tipo de Contrato","Tipo de Movimiento","Calidad Juridica"]
                               if c in df_r.columns]
                        agg = (df_r.groupby(grp, dropna=False)
                               .agg(monto_total_haberes=("_t","sum"), descuentos=("_d","sum"),
                                    haber_neto=("_n","sum"), n_personas=("_u", pd.Series.nunique))
                               .reset_index())

                        rows_imp = [
                            (periodo,
                             str(rw.get("Resolucion","")).strip(), str(rw.get("Programa","")).strip(),
                             str(rw.get("Unidad","")).strip(), str(rw.get("Descripcion Unidad","")).strip(),
                             str(rw.get("Tipo de Contrato","")).strip(),
                             str(rw.get("Tipo de Movimiento","")).strip(), str(rw.get("Calidad Juridica","")).strip(),
                             float(rw["monto_total_haberes"]), float(rw["descuentos"]),
                             float(rw["haber_neto"]), int(rw["n_personas"]), h)
                            for _, rw in agg.iterrows()
                        ]

                        # ── Personas (vectorizado) ─────────────────────────
                        df_p = df_r.copy()
                        df_p["_run"]  = df_p[_cc_run].astype(str).str.strip() if _cc_run else ""
                        df_p["_nom"]  = df_p[(_rename.get(_cc_nom, _cc_nom) if _cc_nom else None)].astype(str).str.strip() if _cc_nom else ""
                        df_p["_res"]  = df_p["Resolucion"].astype(str).str.strip() if "Resolucion" in df_p.columns else ""
                        df_p["_prg"]  = df_p["Programa"].astype(str).str.strip()   if "Programa"   in df_p.columns else ""
                        df_p["_tc"]   = df_p["Tipo de Contrato"].astype(str).str.strip()   if "Tipo de Contrato"   in df_p.columns else ""
                        df_p["_cj"]   = df_p["Calidad Juridica"].astype(str).str.strip()   if "Calidad Juridica"   in df_p.columns else ""
                        df_p["_tm"]   = df_p["Tipo de Movimiento"].astype(str).str.strip()  if "Tipo de Movimiento"  in df_p.columns else ""
                        df_p["_uni"]  = df_p["Unidad"].astype(str).str.strip()              if "Unidad"              in df_p.columns else ""
                        df_p["_duni"] = df_p["Descripcion Unidad"].astype(str).str.strip()  if "Descripcion Unidad"  in df_p.columns else ""
                        df_p["_plp"]  = df_p[(_rename.get(_cc_plp, _cc_plp) if _cc_plp else None)].astype(str).str.strip() if _cc_plp else ""
                        df_p = df_p[~df_p["_run"].isin(["","nan","None","NaN"])]

                        _pcols = ["p_run","p_nom","p_res","p_prg","p_tc","p_cj",
                                  "p_tm","p_uni","p_duni","p_t","p_d","p_n","p_plp"]
                        df_p2 = df_p[["_run","_nom","_res","_prg","_tc","_cj",
                                      "_tm","_uni","_duni","_t","_d","_n","_plp"]].copy()
                        df_p2.columns = _pcols
                        rows_pers = [(h, periodo, r.p_res, r.p_prg, r.p_run, r.p_nom, r.p_tc, r.p_cj,
                                      r.p_tm, r.p_uni, r.p_duni, float(r.p_t), float(r.p_d), float(r.p_n), r.p_plp)
                                     for r in df_p2.itertuples(index=False)]

                        conn = get_conn()
                        _pbar_imp.progress(0.5, text="Insertando imputaciones...")
                        conn.executemany("""INSERT INTO imputaciones
                            (periodo,resolucion,programa,unidad,descripcion_unidad,tipo_contrato,
                             tipo_movimiento,calidad_juridica,monto_total_haberes,descuentos,
                             haber_neto,n_personas,lote_hash) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""", rows_imp)
                        _pbar_imp.progress(0.75, text="Registrando personas...")
                        conn.executemany("""INSERT INTO personas_imputadas
                            (lote_hash,periodo,resolucion,programa,run,nombre,tipo_contrato,
                             calidad_juridica,tipo_movimiento,unidad,descripcion_unidad,
                             monto_total_haberes,descuentos,haber_neto,planilla_pago)
                             VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", rows_pers)
                        conn.execute("INSERT INTO lotes_imputados(lote_hash,periodo,n_registros) VALUES(?,?,?)",(h,periodo,len(df_r)))
                        conn.commit(); conn.close()
                        _pbar_imp.progress(1.0, text="✅ Imputación completada")
                        _imp_status.update(label=f"✅ {per_label(periodo)} imputado", state="complete", expanded=False)
                        st.success(f"✅ **{per_label(periodo)}** imputado — {len(agg):,} grupos · {len(rows_pers):,} personas."); st.rerun()
                    except Exception: st.error("Error al imputar."); st.code(traceback.format_exc())
            except Exception: st.error("Error al leer."); st.code(traceback.format_exc())
    with ch:
        st.markdown("**Historial de imputaciones**")
        lts2 = load_lotes()
        if lts2.empty: st.info("Sin imputaciones.")
        else:
            lts2["Período"] = lts2["periodo"].apply(per_label)
            st.markdown(
                ev_design.ev_table_html(
                    lts2[["Período","n_registros","created_at"]].rename(columns={"n_registros":"Registros","created_at":"Fecha"})
                ),
                unsafe_allow_html=True,
            )
            st.markdown("---")
            sb = st.selectbox("Período a eliminar", lts2["periodo"].tolist(), format_func=per_label, key="sb_del")
            if st.button("🗑️ Eliminar período", use_container_width=True):
                row_d = lts2[lts2["periodo"]==sb].iloc[0]; conn=get_conn()
                conn.execute("DELETE FROM imputaciones WHERE lote_hash=?",(row_d["lote_hash"],))
                conn.execute("DELETE FROM personas_imputadas WHERE lote_hash=?",(row_d["lote_hash"],))
                conn.execute("DELETE FROM lotes_imputados WHERE lote_hash=?",(row_d["lote_hash"],))
                conn.commit(); conn.close(); st.success(f"{per_label(sb)} eliminado."); st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# ESTADO MENSUAL
# ══════════════════════════════════════════════════════════════════════════════
with tab_pivot:
    df_imp = load_imp(); df_cc = load_centros()
    if df_imp.empty:
        st.info("📥 Importa rendiciones para ver el estado mensual.")
    else:
        st.markdown("### 📅 Estado Mensual de Ejecución")
        all_res_p = sorted(df_imp["resolucion"].dropna().unique().tolist())
        f_rp = st.multiselect("Resolución(es)", all_res_p, default=all_res_p, key="piv_r")
        df_p = df_imp[df_imp["resolucion"].isin(f_rp)] if f_rp else df_imp.copy()
        pm = build_prog_map(df_p, df_cc)

        periodos_o = sorted(df_p["periodo"].unique().tolist())
        col_lbl = [per_label(p) for p in periodos_o]

        def mk_piv(df_p, col, periodos_o):
            pv = (df_p.groupby(["resolucion","periodo"])[col].sum().reset_index()
                  .pivot(index="resolucion",columns="periodo",values=col).fillna(0))
            for p in periodos_o:
                if p not in pv.columns: pv[p]=0
            pv = pv[periodos_o]; pv["TOTAL"]=pv.sum(axis=1)
            pv.loc["TOTAL GENERAL"]=pv.sum(); return pv

        # Vista 1: Haber Neto
        st.markdown("#### 💵 Haber Neto por Resolución (CLP)")
        pn = mk_piv(df_p,"haber_neto",periodos_o)
        pn_sorted = pn.drop("TOTAL GENERAL").sort_values("TOTAL",ascending=False)
        pn_sorted.loc["TOTAL GENERAL"] = pn.loc["TOTAL GENERAL"]

        hdr = "".join(f"<th>{c}</th>" for c in col_lbl)+"<th class='tc'>TOTAL</th>"
        rows = ""
        for i,r in pn_sorted.iterrows():
            tg = (i=="TOTAL GENERAL"); cls="tr-tot" if tg else ""
            nom = f"<b>{i}</b>" if tg else str(i)
            prog_txt = "" if tg else f"<div class='tprog'>{pm.get(str(i),'')[:70]}</div>"
            cells = "".join(f"<td>$ {fmt_n(r[p])}</td>" for p in periodos_o)
            cells += f"<td class='tc'>$ {fmt_n(r['TOTAL'])}</td>"
            rows += f"<tr class='{cls}'><td class='tleft'>{nom}{prog_txt}</td>{cells}</tr>"

        st.markdown(f'<div style="overflow-x:auto;margin-bottom:1.2rem">'+
            f'<table class="ptbl"><thead><tr><th class="tleft" style="min-width:220px">Resolución · Programa</th>{hdr}</tr></thead>'+
            f'<tbody>{rows}</tbody></table></div>', unsafe_allow_html=True)

        b1=io.BytesIO(); pn_e=pn.copy(); pn_e.columns=col_lbl+["TOTAL"]; pn_e.to_excel(b1,engine="openpyxl")
        st.download_button("📥 Exportar Haber Neto",b1.getvalue(),f"haber_neto_{datetime.now().strftime('%Y%m%d')}.xlsx",use_container_width=True,key="dl1")

        st.markdown("---")

        # Vista 2: Total + Descuentos + Neto
        st.markdown("#### 📋 Vista Completa: Total Haberes · Descuentos · Haber Neto")
        pt  = mk_piv(df_p,"monto_total_haberes",periodos_o)
        pd2 = mk_piv(df_p,"descuentos",periodos_o)
        pn2 = mk_piv(df_p,"haber_neto",periodos_o)
        pn2_sorted = pn2.drop("TOTAL GENERAL").sort_values("TOTAL",ascending=False)
        all_idx = list(pn2_sorted.index) + ["TOTAL GENERAL"]
        hdr2 = "<th>Concepto</th>"+"".join(f"<th>{c}</th>" for c in col_lbl)+"<th class='tc'>TOTAL</th>"

        def rc(pv,res):
            row = pv.loc[res] if res in pv.index else pd.Series(0,index=periodos_o+["TOTAL"])
            return ("".join(f"<td>$ {fmt_n(row.get(p,0))}</td>" for p in periodos_o)+
                    f"<td class='tc'>$ {fmt_n(row.get('TOTAL',0))}</td>")

        rows2=""
        for res in all_idx:
            tg=(res=="TOTAL GENERAL"); cls="tr-tot" if tg else ""
            b_="<b>" if tg else ""; bc="</b>" if tg else ""
            prog_lbl="" if tg else f"<div class='tprog'>{pm.get(str(res),'')[:60]}</div>"
            rows2+=(f"<tr class='{cls}'>"+
                    f"<td rowspan='3' class='tleft' style='vertical-align:middle;border-right:2px solid #1a3050;min-width:150px'>{b_}{res}{bc}{prog_lbl}</td>"+
                    f"<td class='lth'>Suma Monto (Total Haberes)</td>{rc(pt,res)}</tr>"+
                    f"<tr class='{cls}'><td class='ltd'>Suma Descuentos</td>{rc(pd2,res)}</tr>"+
                    f"<tr class='{cls}' style='border-bottom:2px solid #1a3050'><td class='ltn'>Suma Haber Neto</td>{rc(pn2,res)}</tr>")

        st.markdown(f'<div style="overflow-x:auto;margin-bottom:1.2rem">'+
            f'<table class="ptbl"><thead><tr><th class="tleft" style="min-width:150px">Resolución · Programa</th>{hdr2}</tr></thead>'+
            f'<tbody>{rows2}</tbody></table></div>', unsafe_allow_html=True)

        b2=io.BytesIO()
        with pd.ExcelWriter(b2,engine="openpyxl") as wr:
            pt.rename(columns=dict(zip(periodos_o,col_lbl))).to_excel(wr,sheet_name="Total Haberes")
            pd2.rename(columns=dict(zip(periodos_o,col_lbl))).to_excel(wr,sheet_name="Descuentos")
            pn2.rename(columns=dict(zip(periodos_o,col_lbl))).to_excel(wr,sheet_name="Haber Neto")
        st.download_button("📥 Exportar tabla completa",b2.getvalue(),f"estado_mensual_{datetime.now().strftime('%Y%m%d')}.xlsx",use_container_width=True,key="dl2")

# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
with tab_dash:
    df_imp=load_imp(); df_cc=load_centros()
    if df_imp.empty: st.info("📥 Importa rendiciones para ver el dashboard."); st.stop()
    pd_all=sorted(df_imp["periodo"].unique(),reverse=True)
    td_all=df_imp["tipo_contrato"].dropna().unique().tolist()
    dfc1,dfc2=st.columns([1.5,1.5])
    f_per=dfc1.multiselect("Período(s)",pd_all,default=pd_all,format_func=per_label,key="d_per")
    f_tip=dfc2.multiselect("Tipo contrato",td_all,default=td_all,key="d_tip")
    df_f=df_imp[df_imp["periodo"].isin(f_per)&df_imp["tipo_contrato"].isin(f_tip)]
    if df_f.empty: st.warning("Sin datos."); st.stop()
    pm=build_prog_map(df_f,df_cc)

    tej=df_f["haber_neto"].sum(); thb=df_f["monto_total_haberes"].sum()
    tds=df_f["descuentos"].sum(); tmr=df_cc["marco_clp"].sum() if not df_cc.empty else 0
    pgl=(tej/tmr*100) if tmr>0 else 0
    k1,k2,k3,k4,k5=st.columns(5)
    k1.metric("Haber Neto",fmt_clp(tej)); k2.metric("Total Haberes",fmt_clp(thb))
    k3.metric("Descuentos",fmt_clp(tds)); k4.metric("Marco Presup.",fmt_clp(tmr)); k5.metric("% Ejecución",f"{pgl:.1f}%")
    st.markdown("---")

    st.markdown("### 🏦 Ejecución por Centro de Costo")
    ej=(df_f.groupby("resolucion").agg(ej=("haber_neto","sum"),thb=("monto_total_haberes","sum"),ds=("descuentos","sum")).reset_index().rename(columns={"resolucion":"cc"}))
    ej["cc"]=ej["cc"].astype(str)
    if not df_cc.empty:
        dm=df_cc[["numero","nombre","marco_clp"]].copy(); dm["numero"]=dm["numero"].astype(str)
        ej=ej.merge(dm,left_on="cc",right_on="numero",how="left")
        ej["nombre"]=ej["nombre"].fillna(ej["cc"]); ej["marco_clp"]=ej["marco_clp"].fillna(0)
    else:
        ej["nombre"]=ej["cc"]; ej["marco_clp"]=0
    ej["programa_formal"]=ej["cc"].map(pm).fillna("")
    ej["pct"]=ej.apply(lambda r:r["ej"]/r["marco_clp"]*100 if r["marco_clp"]>0 else 0,axis=1)
    ej=ej.sort_values("ej",ascending=False)

    # ── Lista visual (HTML compacto, igual que antes) ──────────────────────────
    filas_html = []
    for _, row in ej.iterrows():
        nombre       = str(row.get("nombre", row["cc"])).strip()
        tiene_nombre = nombre != str(row["cc"]).strip()
        programa     = str(row.get("programa_formal", "")).strip()
        pct          = row["pct"]
        color_pct    = "#ef4444" if pct >= 95 else "#f59e0b" if pct >= 80 else "#1db954"
        bar_w        = min(pct, 100)
        prog_tag = (f'<div style="font-size:.72rem;color:#6b8caf;margin-top:1px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis">{programa[:70]}</div>') if programa else ""
        nombre_tag = f'<span style="font-weight:700;font-size:.9rem">{nombre if tiene_nombre else ""}</span>'
        badge  = f'<span style="background:#1e3a5f;color:#93c5fd;font-family:monospace;font-size:.75rem;padding:1px 7px;border-radius:4px;margin-right:7px">{row["cc"]}</span>'
        col1   = f'<div>{badge}{nombre_tag}{prog_tag}</div>'
        col2   = f'<div style="text-align:right;padding-right:12px"><div style="font-family:monospace;font-weight:700;font-size:.88rem">{fmt_clp(row["ej"])}</div><div style="font-size:.7rem;color:#6b8caf">Marco: {fmt_clp(row["marco_clp"])}</div></div>'
        barra  = f'<div style="flex:1;background:rgba(255,255,255,0.08);border-radius:4px;height:7px;overflow:hidden"><div style="width:{bar_w}%;height:100%;background:{color_pct};border-radius:4px"></div></div>'
        col3   = f'<div><div style="display:flex;align-items:center;gap:6px">{barra}<span style="color:{color_pct};font-size:.78rem;font-weight:700;white-space:nowrap">{pct:.1f}%</span></div></div>'
        filas_html.append(f'<div style="display:grid;grid-template-columns:3fr 1.4fr 1.8fr;align-items:center;padding:5px 8px;border-bottom:1px solid rgba(255,255,255,0.07)">{col1}{col2}{col3}</div>')

    st.markdown('<div style="border:1px solid rgba(255,255,255,0.1);border-radius:8px;overflow:hidden;margin-bottom:1rem">' + "".join(filas_html) + '</div>', unsafe_allow_html=True)

    # ── Detalle por CC (selector) ──────────────────────────────────────────────
    _cc_opts = {f"[{r['cc']}]  {str(r.get('nombre', r['cc']))[:50]}": str(r["cc"]) for _, r in ej.iterrows()}
    _sel_cc  = st.selectbox("🔎 Ver detalle de Centro de Costo", ["— Selecciona un CC —"] + list(_cc_opts.keys()), key="sel_cc_det")

    if _sel_cc != "— Selecciona un CC —":
        _cc_key  = _cc_opts[_sel_cc]
        _cc_row  = ej[ej["cc"] == _cc_key].iloc[0]
        _cc_data = df_f[df_f["resolucion"].astype(str) == _cc_key]
        _cc_prog = str(_cc_row.get("programa_formal","")).strip()
        _cc_nom  = str(_cc_row.get("nombre", _cc_key)).strip()

        # ── Recuadro contenedor con contraste ──────────────────────────────
        st.markdown(f"""
<div style="border:1px solid #1a73e8;border-radius:12px;padding:20px 24px;
            background:rgba(26,115,232,0.06);margin:12px 0 20px 0">
  <div style="font-size:1.05rem;font-weight:700;color:#93c5fd;margin-bottom:2px">
    <span style="background:#1e3a5f;font-family:monospace;font-size:.8rem;
                 padding:2px 10px;border-radius:4px;margin-right:10px">{_cc_key}</span>
    {_cc_nom}
  </div>
  {"<div style='font-size:.8rem;color:#6b8caf;margin-top:4px'>📌 "+_cc_prog+"</div>" if _cc_prog else ""}
</div>""", unsafe_allow_html=True)

        hd1, hd2, hd3 = st.columns(3)
        hd1.metric("Haber Neto Ejecutado", fmt_clp(_cc_row["ej"]))
        hd2.metric("Marco Presupuestario", fmt_clp(_cc_row["marco_clp"]))
        hd3.metric("% Ejecución del Marco", f"{_cc_row['pct']:.1f}%")

        if not _cc_data.empty:
            # ── Paleta y helpers compartidos ───────────────────────────────
            _MOV_LABEL = {"DESCUENTO_DAP":"Imputación D.A.P.","NETO_ORIGEN":"Gasto Neto Programa"}
            _MOV_CLR   = {"Gasto Neto Programa":"#1a73e8","Gasto Honorarios":"#fbbc05","Imputación D.A.P.":"#ea4335"}
            _CT_CLR    = {"HONORARIOS":"#fbbc05","REMUNERACIONES":"#1a73e8"}

            def _mov_lbl(v):
                s = str(v).strip().upper()
                if s in ("","NAN","NONE","NON","N/A"): return "Gasto Honorarios"
                return _MOV_LABEL.get(s, s)

            def _chip_ct(v):
                s = str(v).strip().upper()
                if "HONOR" in s:
                    return (f'<span style="background:rgba(251,188,5,.15);color:#fbbc05;'
                            f'padding:2px 10px;border-radius:4px;font-weight:700;'
                            f'font-size:10px;letter-spacing:1px;font-family:JetBrains Mono,monospace">{s}</span>')
                return (f'<span style="background:rgba(26,115,232,.15);color:#93c5fd;'
                        f'padding:2px 10px;border-radius:4px;font-weight:700;'
                        f'font-size:10px;letter-spacing:1px;font-family:JetBrains Mono,monospace">{s}</span>')

            _THS = ('padding:10px 14px;font-family:JetBrains Mono,monospace;font-size:9px;'
                    'font-weight:700;letter-spacing:1.8px;text-transform:uppercase;color:#93c5fd;'
                    'background:#1e3a5f;border-bottom:2px solid #1a73e8;white-space:nowrap;')
            _WRAP = ('<div style="background:#181818;border:1px solid rgba(26,115,232,.28);'
                     'border-radius:12px;overflow:auto;box-shadow:0 4px 24px rgba(0,0,0,.5)">'
                     '<table style="width:100%;border-collapse:collapse">')

            # ── preparar datos ─────────────────────────────────────────────
            _t1 = (_cc_data.groupby("tipo_contrato")
                   .agg(haber_neto=("haber_neto","sum"),
                        total_haberes=("monto_total_haberes","sum"),
                        descuentos=("descuentos","sum"),
                        personas=("n_personas","sum"))
                   .reset_index())
            _total_ej = _t1["haber_neto"].sum() or 1
            _t1["pct_v"] = _t1["haber_neto"] / _total_ej * 100

            _t2 = (_cc_data.groupby(["tipo_contrato","tipo_movimiento"])
                   ["haber_neto"].sum().reset_index())
            _t2["_lbl"] = _t2["tipo_movimiento"].apply(_mov_lbl)
            _movs_lbl  = list(dict.fromkeys(_t2["_lbl"].tolist()))
            _contratos = sorted(_t2["tipo_contrato"].unique())

            # ── TABLA 1 + GRÁFICO 1 ────────────────────────────────────────
            st.markdown("---")
            st.markdown("##### Distribución del Gasto por Modalidad Contractual")
            ta1, ta2 = st.columns([1, 1.4], gap="large")

            with ta1:
                _hdr1 = "".join(
                    f'<th style="{_THS}text-align:{"right" if h not in ("Modalidad","% del Total") else "left"}">{h}</th>'
                    for h in ["Modalidad","Haber Neto","Total Haberes","Descuentos","Personas","% del Total"])
                _bdy1 = ""
                for i, (_, r) in enumerate(_t1.iterrows()):
                    _bg = "rgba(255,255,255,.03)" if i%2==0 else "transparent"
                    _pv = r["pct_v"]
                    _bar = (f'<div style="display:flex;align-items:center;gap:6px">'
                            f'<div style="flex:1;background:rgba(255,255,255,.07);border-radius:3px;height:6px">'
                            f'<div style="width:{_pv:.0f}%;height:100%;background:#fbbc05;border-radius:3px"></div></div>'
                            f'<span style="color:#fbbc05;font-size:10px;font-weight:700;'
                            f'font-family:JetBrains Mono,monospace;white-space:nowrap">{_pv:.1f}%</span></div>')
                    _bdy1 += (f'<tr style="background:{_bg}">'
                              f'<td style="padding:9px 14px;border-bottom:1px solid rgba(255,255,255,.05)">{_chip_ct(r["tipo_contrato"])}</td>'
                              f'<td style="padding:9px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#f1f5f9;border-bottom:1px solid rgba(255,255,255,.05)">{fmt_clp(r["haber_neto"])}</td>'
                              f'<td style="padding:9px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;color:#94a3b8;border-bottom:1px solid rgba(255,255,255,.05)">{fmt_clp(r["total_haberes"])}</td>'
                              f'<td style="padding:9px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;color:#94a3b8;border-bottom:1px solid rgba(255,255,255,.05)">{fmt_clp(r["descuentos"])}</td>'
                              f'<td style="padding:9px 14px;text-align:center;font-family:JetBrains Mono,monospace;font-size:11px;color:#94a3b8;border-bottom:1px solid rgba(255,255,255,.05)">{int(r["personas"])}</td>'
                              f'<td style="padding:9px 14px;border-bottom:1px solid rgba(255,255,255,.05)">{_bar}</td>'
                              f'</tr>')
                _bdy1 += (f'<tr style="background:rgba(26,115,232,.18);border-top:2px solid #1a73e8">'
                          f'<td style="padding:10px 14px;font-family:Cabinet Grotesk,sans-serif;font-size:11px;font-weight:800;color:#93c5fd">TOTAL</td>'
                          f'<td style="padding:10px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:12px;font-weight:700;color:#fbbc05">{fmt_clp(_t1["haber_neto"].sum())}</td>'
                          f'<td style="padding:10px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#f1f5f9">{fmt_clp(_t1["total_haberes"].sum())}</td>'
                          f'<td style="padding:10px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#f1f5f9">{fmt_clp(_t1["descuentos"].sum())}</td>'
                          f'<td style="padding:10px 14px;text-align:center;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#f1f5f9">{int(_t1["personas"].sum())}</td>'
                          f'<td style="padding:10px 14px;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#fbbc05">100 %</td>'
                          f'</tr>')
                st.markdown(_WRAP + f'<thead><tr>{_hdr1}</tr></thead><tbody>{_bdy1}</tbody></table></div>', unsafe_allow_html=True)

            with ta2:
                _fc1 = go.Figure()
                for _, _r in _t1.iterrows():
                    _c = _CT_CLR.get(str(_r["tipo_contrato"]).upper(), "#607d8b")
                    _fc1.add_bar(x=[str(_r["tipo_contrato"])], y=[_r["haber_neto"]],
                                 name=str(_r["tipo_contrato"]),
                                 marker_color=_c, marker_line_color="#0f1117", marker_line_width=1.5,
                                 text=[fmt_clp(_r["haber_neto"])], textposition="outside",
                                 textfont=dict(size=12, color="#f1f5f9"))
                _fc1.update_layout(
                    **{**PB, "margin": dict(l=10, r=10, t=55, b=45),
                       "title": dict(text=f"Haber Neto · CC {_cc_key} — {_cc_nom[:28]}",
                                     font=dict(size=13, color="#93c5fd"))},
                    height=320, showlegend=True,
                    legend=dict(orientation="h", y=-0.18, font=dict(size=11), bgcolor="rgba(0,0,0,0)"),
                    yaxis=ax(title="Haber Neto (CLP)"),
                    xaxis=ax(title="Modalidad Contractual"))
                st.plotly_chart(_fc1, use_container_width=True)

            # ── TABLA 2 + GRÁFICO 2 ────────────────────────────────────────
            st.markdown("---")
            st.markdown("##### Descomposición por Tipo de Movimiento Financiero")
            tb1, tb2 = st.columns([1, 1.4], gap="large")

            with tb1:
                _cols2 = ["Modalidad"] + _movs_lbl + ["TOTAL"]
                _hdr2  = "".join(
                    f'<th style="{_THS}text-align:{"right" if c not in ("Modalidad",) else "left"}">{c}</th>'
                    for c in _cols2)
                _piv2 = []
                for _ct in _contratos:
                    _sub = _t2[_t2["tipo_contrato"]==_ct]
                    _r2  = {"Modalidad": _ct}
                    _tot2 = 0
                    for _lb in _movs_lbl:
                        _v = _sub[_sub["_lbl"]==_lb]["haber_neto"].sum()
                        _r2[_lb] = _v; _tot2 += _v
                    _r2["TOTAL"] = _tot2
                    _piv2.append(_r2)
                _tr2 = {"Modalidad":"TOTAL GENERAL"}
                for _lb in _movs_lbl:
                    _tr2[_lb] = _t2[_t2["_lbl"]==_lb]["haber_neto"].sum()
                _tr2["TOTAL"] = _t2["haber_neto"].sum()
                _piv2.append(_tr2)

                _bdy2 = ""
                for i, r2 in enumerate(_piv2):
                    _is_tot2 = str(r2.get("Modalidad","")).strip().upper() == "TOTAL GENERAL"
                    _bg2 = "rgba(26,115,232,.18)" if _is_tot2 else ("rgba(255,255,255,.03)" if i%2==0 else "transparent")
                    _bord2 = "border-top:2px solid #1a73e8;" if _is_tot2 else ""
                    _cells2 = ""
                    for col2 in _cols2:
                        v2 = r2.get(col2, 0)
                        if col2 == "Modalidad":
                            if _is_tot2:
                                _cells2 += f'<td style="padding:10px 14px;font-family:Cabinet Grotesk,sans-serif;font-size:11px;font-weight:800;color:#93c5fd;border-bottom:1px solid rgba(26,115,232,.25)">{v2}</td>'
                            else:
                                _cells2 += f'<td style="padding:9px 14px;border-bottom:1px solid rgba(255,255,255,.05)">{_chip_ct(v2)}</td>'
                        else:
                            _fv2 = fmt_clp(float(v2)) if pd.notna(v2) and float(v2)!=0 else ""
                            _is_tc = col2 == "TOTAL"
                            if _is_tot2:
                                _cc2 = "#fbbc05" if _is_tc else "#f1f5f9"
                                _cells2 += f'<td style="padding:10px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:12px;font-weight:700;color:{_cc2};border-bottom:1px solid rgba(26,115,232,.25)">{_fv2}</td>'
                            elif _is_tc:
                                _cells2 += f'<td style="padding:9px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;font-weight:700;color:#f1f5f9;border-bottom:1px solid rgba(255,255,255,.05)">{_fv2}</td>'
                            else:
                                _cells2 += f'<td style="padding:9px 14px;text-align:right;font-family:JetBrains Mono,monospace;font-size:11px;color:#94a3b8;border-bottom:1px solid rgba(255,255,255,.05)">{_fv2}</td>'
                    _bdy2 += f'<tr style="background:{_bg2};{_bord2}">{_cells2}</tr>'
                st.markdown(_WRAP + f'<thead><tr>{_hdr2}</tr></thead><tbody>{_bdy2}</tbody></table></div>', unsafe_allow_html=True)

            with tb2:
                _PAL2 = ["#1a73e8","#fbbc05","#ea4335","#34a853","#9c27b0","#00bcd4","#ff7043","#607d8b"]
                _fc2  = go.Figure()
                for _i, _lb in enumerate(_movs_lbl):
                    _clr2 = _MOV_CLR.get(_lb, _PAL2[_i % len(_PAL2)])
                    _vals2 = [_t2[(_t2["tipo_contrato"]==_ct)&(_t2["_lbl"]==_lb)]["haber_neto"].sum()
                              for _ct in _contratos]
                    _fc2.add_bar(name=_lb, x=_contratos, y=_vals2,
                                 marker_color=_clr2, marker_line_color="#0f1117", marker_line_width=1,
                                 text=[fmt_clp(v) if v>0 else "" for v in _vals2],
                                 textposition="inside", textfont=dict(size=10, color="#fff"))
                _fc2.update_layout(
                    **{**PB, "margin": dict(l=10, r=10, t=55, b=80),
                       "title": dict(text=f"Estructura de Gasto por Tipo de Movimiento · CC {_cc_key}",
                                     font=dict(size=13, color="#93c5fd"))},
                    height=360, barmode="stack",
                    yaxis=ax(title="Haber Neto (CLP)"),
                    xaxis=ax(title="Modalidad Contractual"),
                    legend=dict(orientation="h", y=-0.32, font=dict(size=11),
                                bgcolor="rgba(0,0,0,0)",
                                bordercolor="rgba(255,255,255,.1)", borderwidth=1))
                st.plotly_chart(_fc2, use_container_width=True)
        else:
            st.info("Sin datos de imputación para este Centro de Costo.")

    # ── CSS: borde sutil para todos los charts de esta sección ──────────────
    st.markdown("""
<style>
section[data-testid="stMain"] div[data-testid="stPlotlyChart"] {
    border: 1px solid rgba(255,255,255,0.13);
    border-radius: 10px;
    padding: 6px 4px 2px 4px;
    background: rgba(255,255,255,0.025);
}
</style>""", unsafe_allow_html=True)

    st.markdown("---")
    cg1, cg2 = st.columns(2, gap="medium")

    with cg1:
        top = ej.head(12).copy()
        top["lbl"] = top.apply(
            lambda r: str(r["nombre"]) if str(r["nombre"]) != str(r["cc"]) else str(r["cc"]), axis=1)
        top["lbl"] = top["lbl"].str[:26]
        f1 = go.Figure()
        if top["marco_clp"].sum() > 0:
            f1.add_bar(
                name="Marco presupuestario", x=top["lbl"], y=top["marco_clp"],
                marker_color="rgba(0,87,255,0.22)",
                marker_line_color="#1a73e8", marker_line_width=1.5)
        colors_bar = ["#ef4444" if p >= 95 else "#f59e0b" if p >= 80 else "#1db954"
                      for p in top["pct"]]
        f1.add_bar(
            name="Haber Neto ejecutado", x=top["lbl"], y=top["ej"],
            marker_color=colors_bar,
            text=[fmt_clp(v) for v in top["ej"]],
            textposition="outside", textfont=dict(size=10, color="#f1f5f9"))
        f1.update_layout(
            **{**PB,
               "title": dict(text="Ejecutado vs Marco Presupuestario · Top CC",
                             font=dict(size=13, color="#93c5fd"),
                             x=0.02, xanchor="left"),
               "margin": dict(l=10, r=10, t=55, b=110)},
            barmode="overlay", height=440,
            xaxis=ax(title="Centro de Costo", tickangle=-45),
            yaxis=ax(title="Haber Neto (CLP)"),
            legend=dict(orientation="h", y=-0.38, font=dict(size=11),
                        bgcolor="rgba(0,0,0,0)"))
        f1.update_xaxes(tickfont=dict(size=9, color="#475569"))
        st.plotly_chart(f1, use_container_width=True)

    with cg2:
        ep = (df_f.groupby("programa")["haber_neto"].sum()
              .reset_index().sort_values("haber_neto", ascending=False).head(10))
        # Paleta categórica de alto contraste — 10 colores distribuidos en el espectro
        PAL10 = ["#1a73e8",  # Azul institucional
                 "#34a853",  # Verde
                 "#fbbc05",  # Ámbar
                 "#ea4335",  # Rojo
                 "#9c27b0",  # Púrpura
                 "#00bcd4",  # Cian
                 "#ff7043",  # Naranja profundo
                 "#607d8b",  # Gris-pizarra
                 "#e91e63",  # Rosa
                 "#795548"]  # Marrón
        ep["lbl_leg"] = ep["programa"].apply(lambda x: str(x)[:48])
        total_pie = ep["haber_neto"].sum() or 1
        ep["pct_v"] = ep["haber_neto"] / total_pie * 100
        pull_vals = [0.06 if i == 0 else 0 for i in range(len(ep))]
        # Solo mostrar % dentro de slices grandes (≥3%), etiquetas en leyenda
        custom_text = [f"{r['pct_v']:.1f}%" if r["pct_v"] >= 3 else ""
                       for _, r in ep.iterrows()]
        f2 = go.Figure(go.Pie(
            labels=ep["lbl_leg"],
            values=ep["haber_neto"],
            hole=0.46,
            pull=pull_vals,
            marker=dict(colors=PAL10[:len(ep)],
                        line=dict(color="#0f1117", width=2)),
            text=custom_text,
            textinfo="text",
            textposition="inside",
            insidetextorientation="auto",
            textfont=dict(size=12, color="#ffffff", family="Arial Black"),
            hovertemplate=(
                "<b>%{label}</b><br>"
                "Haber Neto: $%{value:,.0f}<br>"
                "Participación: %{percent}<extra></extra>"),
            showlegend=True))
        f2.update_layout(
            **{**PB,
               "title": dict(text="Distribución Haber Neto por Programa · Top 10",
                             font=dict(size=13, color="#93c5fd"),
                             x=0.02, xanchor="left"),
               "margin": dict(l=10, r=170, t=55, b=10)},
            height=440,
            legend=dict(
                font=dict(size=10, color="#cbd5e1"),
                orientation="v", x=1.01, y=1.0,
                xanchor="left", yanchor="top",
                bgcolor="rgba(15,17,23,0.7)",
                bordercolor="rgba(255,255,255,0.12)",
                borderwidth=1,
                tracegroupgap=4))
        st.plotly_chart(f2, use_container_width=True)

    st.markdown("### 📈 Evolución Mensual del Gasto")
    ev=(df_f.groupby("periodo").agg(hn=("haber_neto","sum"),thb=("monto_total_haberes","sum"),ds=("descuentos","sum")).reset_index().sort_values("periodo"))
    ev["lbl"]=ev["periodo"].apply(per_label)
    f3=go.Figure()
    f3.add_bar(name="Total Haberes",x=ev["lbl"],y=ev["thb"],
               marker_color="rgba(26,115,232,0.55)",
               marker_line_color="#1a73e8",marker_line_width=1.2)
    f3.add_bar(name="Descuentos",x=ev["lbl"],y=ev["ds"],
               marker_color="rgba(251,188,5,0.60)",
               marker_line_color="#fbbc05",marker_line_width=1.2)
    f3.add_scatter(name="Haber Neto",x=ev["lbl"],y=ev["hn"],mode="lines+markers+text",
                   text=[fmt_clp(v) for v in ev["hn"]],textposition="top center",
                   textfont=dict(size=13,color="#34a853"),
                   line=dict(color="#34a853",width=3),
                   marker=dict(size=12,color="#34a853",line=dict(color="#ffffff",width=1.5)))
    f3.update_layout(
        **{**PB,
           "title": dict(text="Evolución Mensual: Total Haberes · Descuentos · Haber Neto",
                         font=dict(size=13,color="#93c5fd"), x=0.02, xanchor="left"),
           "margin": dict(l=10,r=10,t=55,b=60)},
        barmode="group",height=380,
        xaxis=ax(title="Período"),yaxis=ax(title="CLP"),
        legend=dict(orientation="h",y=-0.22,font=dict(size=13),bgcolor="rgba(0,0,0,0)"))
    st.plotly_chart(f3,use_container_width=True)

    # ══════════════════════════════════════════════════════════════════════
    # BUSCADOR DE PERSONAS
    # ══════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown("### 🔍 Buscador de Personas")

    try:
        _conn_bus = get_conn()
        _df_pers  = pd.read_sql("SELECT * FROM personas_imputadas", _conn_bus)
    finally:
        _conn_bus.close()

    if _df_pers.empty:
        st.info("Sin registros individuales. Re-importa las rendiciones para habilitar el buscador por persona.")
    else:
        _b1, _b2, _b3 = st.columns([2.5, 1.8, 1.5])
        _busq_txt  = _b1.text_input(
            "Nombre o RUT de la persona",
            placeholder="ej: 12345678  ó  JUAN PEREZ",
            key="busq_pers_txt")
        _prog_opts = ["(Todos los programas)"] + sorted(_df_pers["programa"].dropna().unique().tolist())
        _busq_prog = _b2.selectbox("Programa", _prog_opts, key="busq_pers_prog")
        _busq_cc   = _b3.selectbox("CC / Resolución",
                                    ["(Todas)"] + sorted(_df_pers["resolucion"].dropna().unique().tolist()),
                                    key="busq_pers_cc")

        _df_bus = _df_pers.copy()
        if _busq_prog != "(Todos los programas)":
            _df_bus = _df_bus[_df_bus["programa"] == _busq_prog]
        if _busq_cc != "(Todas)":
            _df_bus = _df_bus[_df_bus["resolucion"].astype(str) == str(_busq_cc)]
        if _busq_txt.strip():
            _q = _busq_txt.strip().lower()
            _df_bus = _df_bus[
                _df_bus["run"].astype(str).str.lower().str.contains(_q, na=False) |
                _df_bus["nombre"].astype(str).str.lower().str.contains(_q, na=False)
            ]

        _sor1, _sor2 = st.columns([2, 1])
        _sor1.caption(f"{len(_df_bus):,} registros encontrados")
        _sort_by = _sor2.selectbox(
            "Ordenar por", ["Haber Neto ↓", "Haber Neto ↑", "Total Haberes ↓", "Nombre A→Z"],
            key="busq_sort", label_visibility="collapsed")

        if not _df_bus.empty:
            _sort_map = {
                "Haber Neto ↓":    ("haber_neto", False),
                "Haber Neto ↑":    ("haber_neto", True),
                "Total Haberes ↓": ("monto_total_haberes", False),
                "Nombre A→Z":      ("nombre", True),
            }
            _sc, _sasc = _sort_map.get(_sort_by, ("haber_neto", False))
            _df_show = (_df_bus[["periodo","run","nombre","planilla_pago","resolucion","programa",
                                  "tipo_contrato","calidad_juridica",
                                  "monto_total_haberes","descuentos","haber_neto"]]
                        .copy()
                        .sort_values(_sc, ascending=_sasc)
                        .head(300))
            _df_show["periodo"]             = _df_show["periodo"].apply(per_label)
            _df_show["monto_total_haberes"] = _df_show["monto_total_haberes"].apply(fmt_clp)
            _df_show["descuentos"]          = _df_show["descuentos"].apply(fmt_clp)
            _df_show["haber_neto"]          = _df_show["haber_neto"].apply(fmt_clp)
            _df_show.columns = ["Período","RUT","Nombre","Planilla de Pago","CC","Programa",
                                 "Tipo Contrato","Calidad Jurídica",
                                 "Total Haberes","Descuentos","Haber Neto"]
            st.markdown(ev_design.ev_table_html(_df_show), unsafe_allow_html=True)
            if len(_df_bus) > 300:
                st.caption("⚠️ Se muestran los primeros 300 resultados.")
            _sc1, _sc2, _sc3 = st.columns(3)
            _sc1.metric("Total Haberes", fmt_clp(_df_bus["monto_total_haberes"].sum()))
            _sc2.metric("Descuentos",    fmt_clp(_df_bus["descuentos"].sum()))
            _sc3.metric("Haber Neto",    fmt_clp(_df_bus["haber_neto"].sum()))

# ══════════════════════════════════════════════════════════════════════════════
# ANÁLISIS DETALLADO
# ══════════════════════════════════════════════════════════════════════════════
with tab_det:
    df_imp=load_imp(); df_cc=load_centros()
    if df_imp.empty: st.info("Sin datos. Importa rendiciones primero."); st.stop()
    st.markdown("### 🔍 Análisis de Dotación y Costo")
    pd3=sorted(df_imp["periodo"].unique(),reverse=True)
    da1,da2=st.columns(2)
    fp3=da1.multiselect("Período(s)",pd3,default=pd3,format_func=per_label,key="det_p")
    fr3=da2.multiselect("Resolución(es)",sorted(df_imp["resolucion"].dropna().unique()),
                        default=sorted(df_imp["resolucion"].dropna().unique()),key="det_r")
    df_d=df_imp[df_imp["periodo"].isin(fp3)&df_imp["resolucion"].isin(fr3)]
    if df_d.empty: st.warning("Sin datos."); st.stop()
    pm=build_prog_map(df_d,df_cc)

    # Paleta compartida con gráfico de torta
    _C_HON  = "#fbbc05"   # Ámbar  — Honorarios
    _C_REM  = "#1a73e8"   # Azul   — Remuneraciones
    _C_LINE = "#34a853"   # Verde  — líneas / promedios

    dd1,dd2=st.columns([1,1.1],gap="medium")
    with dd1:
        st.markdown("##### Composición de Dotación y Costo por Modalidad Contractual")
        ta=(df_d.groupby("tipo_contrato").agg(personas=("n_personas","sum"),costo=("haber_neto","sum")).reset_index())
        ct_c={"HONORARIOS": _C_HON, "REMUNERACIONES": _C_REM}
        f4=make_subplots(
            rows=1, cols=2,
            subplot_titles=["N° Personas por Modalidad", "Haber Neto (CLP) por Modalidad"],
            column_widths=[0.48, 0.52])
        for _,r in ta.iterrows():
            c = ct_c.get(str(r["tipo_contrato"]).upper(), "#607d8b")
            f4.add_bar(row=1, col=1,
                       x=[str(r["tipo_contrato"])], y=[r["personas"]],
                       marker_color=c, marker_line_color="#0f1117", marker_line_width=1.5,
                       showlegend=False,
                       text=[fmt_n(r["personas"])], textposition="outside",
                       textfont=dict(size=14, color="#f1f5f9", family="JetBrains Mono, monospace"))
            f4.add_bar(row=1, col=2,
                       x=[str(r["tipo_contrato"])], y=[r["costo"]],
                       marker_color=c, marker_line_color="#0f1117", marker_line_width=1.5,
                       showlegend=False,
                       text=[fmt_clp(r["costo"])], textposition="outside",
                       textfont=dict(size=11, color="#f1f5f9", family="JetBrains Mono, monospace"))
        f4.update_layout(
            **{**PB,
               "title": dict(text="Dotación y Costo Laboral · Honorarios vs Remuneraciones",
                             font=dict(size=13, color="#93c5fd"), x=0.02, xanchor="left"),
               "margin": dict(l=10, r=10, t=70, b=20)},
            height=400, barmode="group",
            xaxis =ax(title="Modalidad"),  yaxis =ax(title="N° Personas"),
            xaxis2=ax(title="Modalidad"),  yaxis2=ax(title="Haber Neto (CLP)"))
        f4.update_annotations(font_size=12, font_color="#93c5fd")
        st.plotly_chart(f4, use_container_width=True)
        ta["Costo CLP"]    = ta["costo"].apply(fmt_clp)
        ta["Prom/persona"] = ta.apply(lambda r: fmt_clp(r["costo"]/r["personas"]) if r["personas"]>0 else "—", axis=1)
        st.markdown(
            ev_design.ev_table_html(
                ta[["tipo_contrato","personas","Costo CLP","Prom/persona"]].rename(
                    columns={"tipo_contrato":"Modalidad","personas":"N° Personas"})),
            unsafe_allow_html=True)

    with dd2:
        st.markdown("##### Top 15 Unidades · Estructura de Costo por Modalidad")
        ua=(df_d.groupby(["descripcion_unidad","tipo_contrato"])
            .agg(costo=("haber_neto","sum")).reset_index()
            .sort_values("costo", ascending=False).head(15))
        f5=px.bar(ua, x="costo", y="descripcion_unidad", color="tipo_contrato",
                  orientation="h",
                  color_discrete_map={"HONORARIOS": _C_HON, "REMUNERACIONES": _C_REM},
                  labels={"costo":"Haber Neto (CLP)", "descripcion_unidad":"Unidad",
                          "tipo_contrato":"Modalidad"},
                  barmode="stack")
        f5.update_traces(marker_line_color="#0f1117", marker_line_width=1)
        f5.update_layout(
            **{**PB,
               "title": dict(text="Haber Neto por Unidad de Desempeño · Top 15",
                             font=dict(size=13, color="#93c5fd"), x=0.02, xanchor="left"),
               "margin": dict(l=10, r=10, t=55, b=50)},
            height=500,
            xaxis=ax(title="Haber Neto (CLP)"), yaxis=ax(title=""),
            legend=dict(orientation="h", y=-0.1, font=dict(size=13),
                        title_text="Modalidad: ", bgcolor="rgba(0,0,0,0)"))
        f5.update_yaxes(autorange="reversed", tickfont=dict(size=11, color="#cbd5e1"))
        st.plotly_chart(f5, use_container_width=True)

    st.markdown("---")
    st.markdown("### 📋 Tabla Dinámica · Haber Neto por Resolución y Tipo de Movimiento")
    st.caption("Estructura equivalente a tabla dinámica Excel. Filas = Resolución · Columnas = Tipo de Movimiento.")

    # ── Mapeo de tipo_movimiento a etiquetas de columna ────────────────────
    _MOV_MAP = {
        "DESCUENTO_DAP": "Imputación a D.A.P. Administración",
        "NETO_ORIGEN":   "Gasto Neto Programa (Haber Neto)",
    }
    _COL_HON  = "Gasto Honorarios"
    _COL_DAP  = "Imputación a D.A.P. Administración"
    _COL_NET  = "Gasto Neto Programa (Haber Neto)"
    _COL_TOT  = "Gasto Total"
    _COL_ORD  = [_COL_DAP, _COL_NET, _COL_HON, _COL_TOT]

    def _mov_col(v):
        s = str(v).strip().upper()
        if s in ("", "NAN", "NONE", "NON", "N/A", "SIN CLASIFICAR"):
            return _COL_HON
        return _MOV_MAP.get(s, s)

    dg = df_d.copy()
    dg["_col"] = dg["tipo_movimiento"].apply(_mov_col)

    # Programa formal por resolución (moda)
    _prog_map = (dg.groupby("resolucion")["programa"]
                 .agg(lambda x: x.mode().iloc[0] if not x.empty else "")
                 .to_dict())

    # Pivot: filas = resolucion, columnas = _col, valores = haber_neto
    _piv = (dg.groupby(["resolucion","_col"])["haber_neto"]
              .sum().unstack(fill_value=0).reset_index())

    for _c in [_COL_DAP, _COL_NET, _COL_HON]:
        if _c not in _piv.columns:
            _piv[_c] = 0.0

    _known = [c for c in [_COL_DAP, _COL_NET, _COL_HON] if c in _piv.columns]
    _piv[_COL_TOT] = _piv[_known].sum(axis=1)

    # Agregar nombre de programa como segunda columna
    _piv.insert(1, "programa", _piv["resolucion"].map(_prog_map).fillna(""))

    # Ordenar numéricamente por resolución
    _piv["_sort"] = pd.to_numeric(_piv["resolucion"], errors="coerce").fillna(9999)
    _piv = _piv.sort_values("_sort").drop(columns=["_sort"])

    # Fila TOTAL GENERAL
    _tot_row = {"resolucion": "TOTAL GENERAL", "programa": ""}
    for _c in _known + [_COL_TOT]:
        _tot_row[_c] = _piv[_c].sum()
    _piv = pd.concat([_piv, pd.DataFrame([_tot_row])], ignore_index=True)

    _extra = [c for c in _piv.columns
              if c not in ["resolucion","programa"] + _COL_ORD and not c.startswith("_")]
    _money_cols = [c for c in _COL_ORD if c in _piv.columns] + _extra
    _final_cols = ["resolucion","programa"] + _money_cols
    _piv = _piv[_final_cols]

    # ── Renderer HTML experto para esta tabla ──────────────────────────────
    def _pivot_table_html(df):
        _MONEY = set(_money_cols)
        # cabecera
        th = ""
        for col in df.columns:
            is_m = col in _MONEY
            align = "right" if is_m else "left"
            th += (f'<th style="padding:11px 16px;text-align:{align};'
                   f'font-family:JetBrains Mono,monospace;font-size:9.5px;font-weight:700;'
                   f'letter-spacing:1.8px;text-transform:uppercase;color:#93c5fd;'
                   f'background:#1e3a5f;border-bottom:2px solid #1a73e8;'
                   f'white-space:nowrap;">{col}</th>')
        # filas
        body = ""
        for i, (_, row) in enumerate(df.iterrows()):
            is_tot = str(row.get("resolucion","")).strip().upper() == "TOTAL GENERAL"
            if is_tot:
                row_bg = "background:rgba(26,115,232,0.2);border-top:2px solid #1a73e8;"
            else:
                row_bg = "background:rgba(255,255,255,0.028);" if i%2==0 else "background:transparent;"
            cells = ""
            for col in df.columns:
                val = str(row[col]) if pd.notna(row[col]) else ""
                is_m   = col in _MONEY
                is_res = col == "resolucion"
                is_tot_col = col == _COL_TOT
                if is_tot:
                    if is_m:
                        color = "#fbbc05" if is_tot_col else "#f1f5f9"
                        cells += (f'<td style="padding:11px 16px;text-align:right;'
                                  f'font-family:JetBrains Mono,monospace;font-size:12px;'
                                  f'font-weight:700;color:{color};letter-spacing:.3px;'
                                  f'border-bottom:1px solid rgba(26,115,232,0.25);">{val}</td>')
                    else:
                        txt = "TOTAL GENERAL" if is_res else val
                        cells += (f'<td style="padding:11px 16px;text-align:left;'
                                  f'font-family:Cabinet Grotesk,sans-serif;font-size:12px;'
                                  f'font-weight:800;color:#93c5fd;letter-spacing:.5px;'
                                  f'border-bottom:1px solid rgba(26,115,232,0.25);">{txt}</td>')
                elif is_res:
                    cells += (f'<td style="padding:9px 16px;text-align:left;'
                              f'font-family:JetBrains Mono,monospace;font-size:11px;'
                              f'font-weight:600;color:#93c5fd;border-bottom:1px solid rgba(255,255,255,.05);">'
                              f'<span style="background:#1e3a5f;padding:2px 9px;border-radius:5px;">{val}</span></td>')
                elif is_m:
                    color = "#f1f5f9" if is_tot_col else "#94a3b8"
                    fw    = "600" if is_tot_col else "400"
                    cells += (f'<td style="padding:9px 16px;text-align:right;'
                              f'font-family:JetBrains Mono,monospace;font-size:11px;'
                              f'font-weight:{fw};color:{color};'
                              f'border-bottom:1px solid rgba(255,255,255,.05);">{val}</td>')
                else:
                    cells += (f'<td style="padding:9px 16px;text-align:left;'
                              f'font-family:Outfit,sans-serif;font-size:11px;'
                              f'color:#b3b3b3;border-bottom:1px solid rgba(255,255,255,.05);">{val}</td>')
            body += f'<tr style="{row_bg}">{cells}</tr>'

        return (f'<div style="background:#181818;border:1px solid rgba(26,115,232,0.28);'
                f'border-radius:12px;overflow:auto;margin-top:8px;'
                f'box-shadow:0 6px 32px rgba(0,0,0,.55);">'
                f'<table style="width:100%;border-collapse:collapse;min-width:820px;">'
                f'<thead><tr>{th}</tr></thead>'
                f'<tbody>{body}</tbody>'
                f'</table></div>')

    # Formatear montos antes de renderizar
    _piv_disp = _piv.copy()
    for _c in _money_cols:
        _piv_disp[_c] = _piv_disp[_c].apply(
            lambda v: fmt_clp(v) if pd.notna(v) and v != 0 else "")
    _piv_disp = _piv_disp.rename(columns={"resolucion":"Resolución","programa":"Programa"})

    st.markdown(_pivot_table_html(_piv_disp), unsafe_allow_html=True)

    # Mantener dg para la exportación Excel
    dg = (df_d.groupby(["resolucion","programa","tipo_movimiento","tipo_contrato"])
          .agg(personas=("n_personas","sum"), monto_total_haberes=("monto_total_haberes","sum"),
               descuentos=("descuentos","sum"), haber_neto=("haber_neto","sum"))
          .reset_index())

    st.markdown("---")
    st.markdown("### 👥 Costo por Calidad Jurídica (homologado)")
    cj1c,cj2c=st.columns(2,gap="medium")
    with cj1c:
        cj=(df_d.groupby("calidad_juridica").agg(personas=("n_personas","sum"),costo=("haber_neto","sum")).reset_index().sort_values("costo",ascending=False))
        cj["prom"]=cj.apply(lambda r:r["costo"]/r["personas"] if r["personas"]>0 else 0,axis=1)
        # PAL10 — mismo sistema de color del gráfico de torta
        _PAL10=["#1a73e8","#34a853","#fbbc05","#ea4335","#9c27b0",
                "#00bcd4","#ff7043","#607d8b","#e91e63","#795548"]
        _cols_cj = [_PAL10[i % len(_PAL10)] for i in range(len(cj))]
        f6=make_subplots(specs=[[{"secondary_y":True}]])
        f6.add_bar(x=cj["calidad_juridica"], y=cj["costo"],
                   name="Haber Neto Total",
                   marker_color=_cols_cj,
                   marker_line_color="#0f1117", marker_line_width=1.5,
                   secondary_y=False,
                   text=[fmt_clp(v) for v in cj["costo"]],
                   textposition="outside",
                   textfont=dict(size=11, color="#f1f5f9", family="JetBrains Mono, monospace"))
        f6.add_scatter(x=cj["calidad_juridica"], y=cj["prom"],
                       name="Costo Promedio / Persona",
                       mode="lines+markers+text",
                       text=[fmt_clp(v) for v in cj["prom"]],
                       textposition="top center",
                       textfont=dict(size=10, color=_C_LINE),
                       marker=dict(size=11, color=_C_LINE,
                                   line=dict(color="#ffffff", width=1.5)),
                       line=dict(color=_C_LINE, dash="dot", width=2.5),
                       secondary_y=True)
        f6.update_layout(
            **{**PB,
               "title": dict(text="Haber Neto Total y Costo Promedio por Persona · Calidad Jurídica",
                             font=dict(size=13, color="#93c5fd"), x=0.02, xanchor="left"),
               "margin": dict(l=10, r=10, t=55, b=60)},
            height=420,
            legend=dict(orientation="h", y=-0.2, font=dict(size=12), bgcolor="rgba(0,0,0,0)"))
        f6.update_xaxes(gridcolor="rgba(255,255,255,.06)",
                        tickfont=dict(size=11, color="#cbd5e1"))
        f6.update_yaxes(title_text="Haber Neto (CLP)",
                        gridcolor="rgba(255,255,255,.06)",
                        tickfont=dict(size=11), secondary_y=False)
        f6.update_yaxes(title_text="Promedio CLP / Persona",
                        gridcolor="rgba(0,0,0,0)",
                        tickfont=dict(size=11), secondary_y=True)
        st.plotly_chart(f6, use_container_width=True)
    with cj2c:
        cj["Total CLP"]=cj["costo"].apply(fmt_clp); cj["Prom/persona CLP"]=cj["prom"].apply(fmt_clp)
        st.markdown(
            ev_design.ev_table_html(
                cj[["calidad_juridica","personas","Total CLP","Prom/persona CLP"]].rename(columns={"calidad_juridica":"Calidad Jurídica","personas":"Personas"})
            ),
            unsafe_allow_html=True,
        )

    st.markdown("---")
    _exp1, _exp2 = st.columns(2, gap="medium")

    with _exp1:
        if st.button("📊 Exportar análisis detallado (Excel)", use_container_width=True):
            bd = io.BytesIO()
            with pd.ExcelWriter(bd, engine="openpyxl") as wr:
                dg.to_excel(wr, sheet_name="Descomposición", index=False)
                ta.to_excel(wr, sheet_name="Tipo Contrato", index=False)
                cj[["calidad_juridica","personas","costo","prom"]].to_excel(wr, sheet_name="Calidad Jurídica", index=False)
            st.download_button("⬇️ Descargar Excel", bd.getvalue(),
                               f"analisis_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               use_container_width=True, key="dl_det")

    with _exp2:
        st.markdown("**Reporte Ejecutivo Word — Subtítulo 21**")
        _periodos_word = sorted(load_lotes()["periodo"].tolist(), reverse=True)
        if len(_periodos_word) < 2:
            st.info("Se necesitan al menos 2 períodos importados para generar el informe comparativo.")
        else:
            _w1, _w2 = st.columns(2)
            _per_act = _w1.selectbox("Período actual",  _periodos_word,
                                     index=0, format_func=per_label, key="word_act")
            _per_ant_opts = [p for p in _periodos_word if p != _per_act]
            _per_ant = _w2.selectbox("Período anterior", _per_ant_opts,
                                     index=0, format_func=per_label, key="word_ant")
            if st.button("📝 Generar Informe Word", use_container_width=True, type="primary", key="btn_word"):
                with st.status("📝 Generando Informe Ejecutivo Word...", expanded=True) as _word_status:
                    _pbar_w  = st.progress(0, text="Iniciando generación del informe...")
                    _step_lbl = st.empty()
                    def _word_cb(step, total, label):
                        _pbar_w.progress(min(step / total, 1.0), text=label)
                        _step_lbl.caption(f"Paso {step}/{total} — {label}")
                    try:
                        _word_buf = generar_informe_word(_per_act, _per_ant, on_progress=_word_cb)
                        _pbar_w.progress(1.0, text="✅ Informe listo")
                        _step_lbl.empty()
                        _word_status.update(label="✅ Informe generado correctamente", state="complete", expanded=False)
                        _fname = f"Reporte_Sub21_DAP_{_per_act}_{datetime.now().strftime('%Y%m%d')}.docx"
                        st.download_button(
                            "⬇️ Descargar Informe Word",
                            _word_buf.getvalue(),
                            file_name=_fname,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="dl_word")
                        st.success("✅ Informe generado. Haz clic en el botón para descargarlo.")
                    except Exception:
                        _word_status.update(label="❌ Error al generar el informe", state="error", expanded=True)
                        st.error("Error al generar el informe.")
                        st.code(traceback.format_exc())
